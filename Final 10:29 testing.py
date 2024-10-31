import warnings

warnings.filterwarnings("ignore", category=UserWarning, module="resource_tracker")

import logging
import os
import subprocess
import sys
import threading
import time
from importlib import metadata
from typing import List, Dict, Optional, Callable, Tuple
import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox
import queue
import datetime
import re
from typing import Dict, Any

# Add type ignore comments for packages that will be installed later
try:
    import numpy as np  # type: ignore
    import spacy  # type: ignore
    import torch  # type: ignore
    from transformers import (  # type: ignore
        AutoTokenizer,
        AutoModelForTokenClassification,
        AutoModelForSequenceClassification,
        pipeline,
    )
except ImportError:
    pass  # Will be handled during dependency setup


# Path configurations
PYTHON312_PATH = "/opt/homebrew/opt/python@3.12/bin/python3.12"
VENV_DIR = "myenv"
VENV_PYTHON_PATH = os.path.join(VENV_DIR, "bin", "python")

# Rate limiting configurations
REQUESTS_PER_MINUTE = 500  # Maximum requests per minute
REQUEST_INTERVAL = 60.0 / REQUESTS_PER_MINUTE  # Time between requests in seconds
MAX_RETRIES = 3  # Maximum number of retries for rate-limited requests
RETRY_DELAY = 30  # Delay in seconds before retrying after rate limit

# Configure logging
logging.basicConfig(
    filename="qa_extractor.log",
    level=logging.INFO,
    format="%(asctime)s %(levelname)s:%(message)s",
)


class RateLimiter:
    """Rate limiter to control API request frequency"""

    def __init__(self, requests_per_minute: int):
        self.interval = 60.0 / requests_per_minute
        self.last_request = 0
        self.lock = threading.Lock()

    def wait(self) -> None:
        """Wait if necessary to comply with rate limit"""
        with self.lock:
            current_time = time.time()
            time_since_last = current_time - self.last_request
            if time_since_last < self.interval:
                sleep_time = self.interval - time_since_last
                time.sleep(sleep_time)
            self.last_request = time.time()


def install_missing_packages(packages: List[str]) -> None:
    """Install missing packages using pip"""
    if not packages:
        return

    logging.info("Installing missing packages: %s", packages)
    print("Installing missing packages:")
    for package in packages:
        print(f"  - {package}")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
            logging.info("Successfully installed package: %s", package)
        except subprocess.CalledProcessError as e:
            logging.error("Failed to install package %s: %s", package, e, exc_info=True)
            messagebox.showerror(
                "Installation Error",
                f"Failed to install package '{package}'.\nError: {e}",
            )
            sys.exit(1)
        except Exception as e:
            logging.error(
                "Unexpected error installing package %s: %s", package, e, exc_info=True
            )
            messagebox.showerror(
                "Installation Error",
                f"An unexpected error occurred while installing '{package}'.\nError: {e}",
            )
            sys.exit(1)


def check_package_version(package_spec: str) -> bool:
    """Check if a package is installed and meets version requirements"""
    try:
        if ">=" in package_spec:
            pkg_name, version_required = package_spec.split(">=")
            pkg_name = pkg_name.strip()
            version_required = version_required.strip()
        else:
            pkg_name, version_required = package_spec.strip(), None

        try:
            installed_version = metadata.version(pkg_name)
            logging.info(
                "Package '%s' version '%s' is installed.", pkg_name, installed_version
            )
        except metadata.PackageNotFoundError:
            logging.warning("Package '%s' is not installed.", pkg_name)
            return False

        if not version_required:
            return True

        from packaging import version

        if version.parse(installed_version) >= version.parse(version_required):
            return True
        else:
            logging.warning(
                "Package '%s' version '%s' does not meet the required version '%s'.",
                pkg_name,
                installed_version,
                version_required,
            )
            return False

    except ValueError:
        logging.error("Invalid package specification: '%s'", package_spec)
        return False
    except Exception as e:
        logging.error(
            "Error checking package version for '%s': %s",
            package_spec,
            e,
            exc_info=True,
        )
        return False


def get_missing_packages(required_packages: List[str]) -> List[str]:
    """Get list of required packages that are not installed or don't meet version requirements"""
    missing_packages = []
    for package in required_packages:
        if not check_package_version(package):
            missing_packages.append(package)
    logging.info("Missing packages: %s", missing_packages)
    return missing_packages


def setup_dependencies() -> None:
    """Check and install all required dependencies"""
    # First ensure packaging and core dependencies are installed
    core_packages = ["packaging", "spacy", "torch", "transformers"]
    for package in core_packages:
        try:
            __import__(package)
            logging.info(f"Package '{package}' is already installed.")
        except ImportError:
            logging.info(f"Package '{package}' not found. Attempting to install.")
            try:
                subprocess.check_call([sys.executable, "-m", "pip", "install", package])
                logging.info(f"Successfully installed '{package}'.")
            except subprocess.CalledProcessError as e:
                logging.error(f"Failed to install '{package}': %s", e, exc_info=True)
                messagebox.showerror(
                    "Dependency Error",
                    f"Failed to install '{package}'.\nError: {e}",
                )
                sys.exit(1)
            except Exception as e:
                logging.error(
                    f"Unexpected error installing '{package}': %s", e, exc_info=True
                )
                messagebox.showerror(
                    "Dependency Error",
                    f"An unexpected error occurred while installing '{package}'.\nError: {e}",
                )
                sys.exit(1)

    required_packages = [
        "openai>=1.0.0",
        "pandas",
        "nltk",
        "spacy>=3.0.0",
        "python-docx",
        "openpyxl",
        "packaging",
        "transformers",
        "torch",
        "scipy",
        "sentencepiece",
        "numpy",
    ]

    logging.info("Checking package dependencies...")
    print("Checking package dependencies...")
    missing_packages = get_missing_packages(required_packages)

    # Manually add numpy to the missing packages list if it's not installed
    try:
        import numpy
    except ImportError:
        logging.info("Forcing installation of numpy.")
        missing_packages.append("numpy")

    if missing_packages:
        logging.info("Missing required packages: %s", missing_packages)
        print("\nMissing required packages:")
        for package in missing_packages:
            print(f"  - {package}")

        install_missing_packages(missing_packages)
        print("\nAll required packages installed successfully.")
        logging.info("All required packages installed successfully.")
    else:
        print("All required packages are already installed.")
        logging.info("All required packages are already installed.")

    if missing_packages:
        logging.info("Missing required packages: %s", missing_packages)
        print("\nMissing required packages:")
        for package in missing_packages:
            print(f"  - {package}")

        try:
            install_missing_packages(missing_packages)
            print("\nAll required packages installed successfully.")
            logging.info("All required packages installed successfully.")
        except Exception as e:
            logging.error("Error installing packages: %s", e, exc_info=True)
            messagebox.showerror(
                "Installation Error",
                f"An error occurred while installing packages.\nError: {e}",
            )
            sys.exit(1)
    else:
        print("All required packages are already installed.")
        logging.info("All required packages are already installed.")

    # Download spaCy model
    try:
        logging.info("Checking spaCy language model...")
        print("Checking spaCy language model...")
        import spacy

        try:
            spacy.load("en_core_web_lg")
            print("spaCy language model already installed.")
            logging.info("spaCy language model already installed.")
        except OSError:
            print("Downloading spaCy language model (this may take a while)...")
            logging.info("Downloading spaCy language model...")
            subprocess.check_call(
                [sys.executable, "-m", "spacy", "download", "en_core_web_lg"]
            )
            print("spaCy language model installed successfully.")
            logging.info("spaCy language model installed successfully.")
    except Exception as e:
        logging.error("Error setting up spaCy model: %s", e, exc_info=True)
        messagebox.showerror(
            "Dependency Error",
            f"Failed to set up spaCy language model.\nError: {e}",
        )
        sys.exit(1)


def is_python312() -> bool:
    """Check if running Python 3.12"""
    return sys.version_info.major == 3 and sys.version_info.minor == 12


def reexecute_with_python312() -> None:
    """Re-run the script using Python 3.12"""
    try:
        logging.info("Re-running the script using Python 3.12 at %s", PYTHON312_PATH)
        subprocess.check_call([PYTHON312_PATH] + sys.argv)
        sys.exit()
    except FileNotFoundError:
        logging.error("Python 3.12 executable not found at %s", PYTHON312_PATH)
        messagebox.showerror(
            "Python Version Error",
            f"Python 3.12 executable not found at {PYTHON312_PATH}. Please install Python 3.12.",
        )
        sys.exit(1)
    except Exception as e:
        logging.error(
            "Failed to re-execute script with Python 3.12: %s", e, exc_info=True
        )
        messagebox.showerror(
            "Execution Error",
            f"Failed to re-execute script with Python 3.12.\nError: {e}",
        )
        sys.exit(1)


def is_venv_active() -> bool:
    """Check if running in a virtual environment"""
    return hasattr(sys, "real_prefix") or (
        hasattr(sys, "base_prefix") and sys.base_prefix != sys.prefix
    )


def create_virtual_env(venv_dir: str) -> None:
    """Create a new virtual environment"""
    try:
        logging.info("Creating virtual environment in %s using Python 3.12", venv_dir)
        print(f"Creating virtual environment in {venv_dir} using Python 3.12...")
        subprocess.check_call([PYTHON312_PATH, "-m", "venv", venv_dir])
        logging.info("Virtual environment created successfully.")
        print("Virtual environment created.")
    except subprocess.CalledProcessError as e:
        logging.error("Failed to create virtual environment: %s", e, exc_info=True)
        messagebox.showerror(
            "Virtual Environment Error",
            f"Failed to create virtual environment.\nError: {e}",
        )
        sys.exit(1)
    except Exception as e:
        logging.error(
            "Unexpected error creating virtual environment: %s", e, exc_info=True
        )
        messagebox.showerror(
            "Virtual Environment Error",
            f"An unexpected error occurred while creating virtual environment.\nError: {e}",
        )
        sys.exit(1)


def reexecute_with_venv(venv_dir: str) -> None:
    """Re-run the script within the virtual environment"""
    python_executable = os.path.join(venv_dir, "bin", "python")
    try:
        logging.info(
            "Re-running the script inside the virtual environment using %s",
            python_executable,
        )
        print(
            f"Re-running the script inside the virtual environment using {python_executable}..."
        )
        subprocess.check_call([python_executable] + sys.argv)
        sys.exit()
    except FileNotFoundError:
        logging.error(
            "Virtual environment Python executable not found at %s", python_executable
        )
        messagebox.showerror(
            "Virtual Environment Error",
            f"Python executable not found in virtual environment at {python_executable}.",
        )
        sys.exit(1)
    except subprocess.CalledProcessError as e:
        logging.error(
            "Failed to execute script in virtual environment: %s", e, exc_info=True
        )
        messagebox.showerror(
            "Execution Error",
            f"Failed to execute script in virtual environment.\nError: {e}",
        )
        sys.exit(1)
    except Exception as e:
        logging.error(
            "Unexpected error executing in virtual environment: %s", e, exc_info=True
        )
        messagebox.showerror(
            "Execution Error",
            f"An unexpected error occurred while executing in virtual environment.\nError: {e}",
        )
        sys.exit(1)


class TranscriptProcessor:
    """Enhanced transcript processor with advanced NLP capabilities for analysis and anonymization"""

    def __init__(self):
        """Enhanced transcript processor with advanced NLP capabilities for analysis and anonymization"""
        self.topic_analyzer = None  # Initialize the topic_analyzer attribute here
        try:
            import spacy
            from transformers import (
                AutoTokenizer,
                AutoModelForTokenClassification,
                AutoModelForSequenceClassification,
            )
            from transformers import pipeline
            import torch

            logging.info("Loading spaCy model...")
            self.spacy_model = spacy.load("en_core_web_lg")

            logging.info("Loading NER model...")
            # Initialize BERT-based NER model
            self.tokenizer = AutoTokenizer.from_pretrained(
                "Jean-Baptiste/roberta-large-ner-english"
            )
            self.model = AutoModelForTokenClassification.from_pretrained(
                "Jean-Baptiste/roberta-large-ner-english"
            )
            self.ner_pipeline = pipeline(
                "ner",
                model=self.model,
                tokenizer=self.tokenizer,
                aggregation_strategy="simple",
                device=0 if torch.cuda.is_available() else -1,  # Use GPU if available
            )

            logging.info(
                "Loading sentiment analysis model (this may take a while on first run)..."
            )
            self.sentiment_tokenizer = AutoTokenizer.from_pretrained(
                "cardiffnlp/twitter-roberta-base-sentiment"
            )
            self.sentiment_model = AutoModelForSequenceClassification.from_pretrained(
                "cardiffnlp/twitter-roberta-base-sentiment"
            )
            self.sentiment_pipeline = pipeline(
                "sentiment-analysis",
                model=self.sentiment_model,
                tokenizer=self.sentiment_tokenizer,
                device=0 if torch.cuda.is_available() else -1,
            )
            logging.info("Loading emotion detection model...")
            self.emotion_tokenizer = AutoTokenizer.from_pretrained(
                "j-hartmann/emotion-english-distilroberta-base"
            )
            self.emotion_model = AutoModelForSequenceClassification.from_pretrained(
                "j-hartmann/emotion-english-distilroberta-base"
            )
            self.emotion_pipeline = pipeline(
                "text-classification",
                model=self.emotion_model,
                tokenizer=self.emotion_tokenizer,
                device=0 if torch.cuda.is_available() else -1,
                top_k=None,
            )
            logging.info("Emotion detection model loaded successfully")
            # Entity type mappings
            self.entity_type_map = {
                # SpaCy entities
                "PERSON": "PERSON",
                "ORG": "ORGANIZATION",
                "GPE": "LOCATION",
                "LOC": "LOCATION",
                "DATE": "DATE",
                "TIME": "TIME",
                "MONEY": "MONEY",
                "PRODUCT": "PRODUCT",
                # RoBERTa entities
                "PER": "PERSON",
                "ORG": "ORGANIZATION",
                "LOC": "LOCATION",
                "MISC": "MISC",
                # Additional context-specific entities
                "ROLE": "ROLE",
                "PROJECT": "PROJECT",
                "SYSTEM": "SYSTEM",
            }

            # Initialize grammatical role mappings
            self.grammatical_roles = {
                "nsubj": "SUBJECT",  # Nominal subject
                "dobj": "OBJECT",  # Direct object
                "pobj": "OBJECT",  # Object of preposition
                "iobj": "INDIRECT_OBJECT",  # Indirect object
                "compound": "COMPOUND",  # Compound name components
                "appos": "APPOSITIVE",  # Appositive (explanatory) phrases
            }

            # Initialize other attributes
            self.client = None
            self.model_name = "gpt-4o"
            self.max_tokens = 8000
            self.rate_limiter = RateLimiter(REQUESTS_PER_MINUTE)
            self.retry_count = 0

            logging.info("Successfully initialized all NLP models and components.")
        except Exception as e:
            logging.error("Error initializing NLP models: %s", e, exc_info=True)
            messagebox.showerror(
                "NLP Model Error",
                f"Failed to initialize NLP models.\nError: {e}",
            )
            sys.exit(1)

    def init_topic_analyzer(self):
        """Lazy initialization of topic analyzer"""
        if self.topic_analyzer is None:
            if not hasattr(self, "_topic_dependencies_installed"):
                setup_topic_modeling_dependencies()
                self._topic_dependencies_installed = True
            self.topic_analyzer = TopicAnalyzer()

    def analyze_answer_sentiment(self, answer_text: str) -> Dict:
        """Analyze the emotional tone of text with natural language output"""
        try:
            # Get basic sentiment
            sentiment_result = self.sentiment_pipeline(answer_text)[0]

            # Map basic sentiment to simple terms
            sentiment_mapping = {
                "LABEL_0": "negative",
                "LABEL_1": "neutral",
                "LABEL_2": "positive",
            }
            overall_sentiment = sentiment_mapping.get(
                sentiment_result["label"], "neutral"
            )

            try:
                # Get detailed emotions
                emotions_result = self.emotion_pipeline(answer_text)

                # Map emotion labels to more natural terms
                emotion_mapping = {
                    "anger": "angry",
                    "disgust": "frustrated",
                    "fear": "concerned",
                    "joy": "enthusiastic",
                    "neutral": "neutral",
                    "sadness": "disappointed",
                    "surprise": "surprised",
                }

                # Process emotions and their intensities
                emotions = []
                for emotion in emotions_result[0]:
                    mapped_label = emotion_mapping.get(
                        emotion["label"], emotion["label"]
                    )
                    emotions.append({"label": mapped_label, "score": emotion["score"]})

                # Sort by intensity
                sorted_emotions = sorted(
                    emotions, key=lambda x: x["score"], reverse=True
                )

                # Generate natural language summary
                summary = self._generate_emotion_summary(
                    overall_sentiment, sorted_emotions
                )

                return {
                    "sentiment_summary": summary,
                    "details": {
                        "overall_sentiment": overall_sentiment,
                        "emotions": sorted_emotions,
                    },
                }

            except Exception as e:
                logging.error(f"Error in emotion detection: {str(e)}")
                # Fallback to basic sentiment
                return {
                    "sentiment_summary": self._get_basic_sentiment_description(
                        overall_sentiment
                    ),
                    "details": {"overall_sentiment": overall_sentiment, "emotions": []},
                }

        except Exception as e:
            logging.error("Error in emotion analysis: %s", e, exc_info=True)
            return {
                "sentiment_summary": "Unable to analyze emotional tone",
                "details": {},
            }

    def _get_basic_sentiment_description(self, sentiment: str) -> str:
        """Convert basic sentiment to natural language"""
        if sentiment == "positive":
            return "The response has a positive tone"
        elif sentiment == "negative":
            return "The response has a negative tone"
        return "The response has a neutral tone"

    def _generate_emotion_summary(self, sentiment: str, emotions: List[Dict]) -> str:
        """Generate a natural-sounding emotional summary"""
        try:
            if not emotions:
                return self._get_basic_sentiment_description(sentiment)

            # Get primary emotion
            primary = emotions[0]
            primary_score = primary["score"]

            # Get strong secondary emotions (threshold lowered for more natural descriptions)
            secondary = [e for e in emotions[1:] if e["score"] > 0.15]

            # Build natural language description
            if primary_score > 0.6:  # Strong primary emotion
                description = f"The response is clearly {primary['label']}"
            elif primary_score > 0.3:  # Moderate primary emotion
                description = f"The response sounds {primary['label']}"
            else:  # Weak emotions
                return f"The response has a {sentiment} tone"

            # Add secondary emotions if present
            if secondary:
                if len(secondary) == 1:
                    description += f" and somewhat {secondary[0]['label']}"
                elif len(secondary) == 2:
                    description += f", with hints of being {secondary[0]['label']} and {secondary[1]['label']}"

            return description

        except Exception as e:
            logging.error("Error generating emotion summary: %s", e, exc_info=True)
            return self._get_basic_sentiment_description(sentiment)

    def load_custom_replacements(self, file_path: str) -> Dict[str, str]:
        """Load custom word replacements from CSV file with case-insensitive handling"""
        import pandas as pd

        try:
            df = pd.read_csv(file_path)
            if len(df.columns) < 2:
                raise ValueError(
                    "CSV must have at least 2 columns: word to replace and replacement"
                )
            replacements = dict(zip(df.iloc[:, 0].str.lower(), df.iloc[:, 1]))
            logging.info(
                "Loaded %d custom replacements from '%s'", len(replacements), file_path
            )
            return replacements
        except Exception as e:
            logging.error(
                "Error loading custom replacements from '%s': %s",
                file_path,
                e,
                exc_info=True,
            )
            raise

    def handle_rate_limit(self, e: Exception) -> bool:
        """Handle rate limit errors with exponential backoff"""
        if "rate_limit" in str(e).lower():
            if self.retry_count < MAX_RETRIES:
                self.retry_count += 1
                wait_time = RETRY_DELAY * (2 ** (self.retry_count - 1))
                logging.warning(
                    f"Rate limited. Waiting {wait_time} seconds before retry {self.retry_count}"
                )
                time.sleep(wait_time)
                return True
            else:
                logging.error("Max retries reached for rate limit")
                return False
        return False

    def set_api_key(self, api_key: str) -> None:
        """Set OpenAI API key"""
        try:
            import openai

            self.client = openai.Client(api_key=api_key)
            logging.info("OpenAI API key set successfully.")
        except Exception as e:
            logging.error("Error setting OpenAI API key: %s", e, exc_info=True)
            raise

    def load_questions(self, file_path: str) -> List[str]:
        """Load questions from Excel file"""
        import pandas as pd

        try:
            if not os.path.exists(file_path):
                logging.error("Questions file '%s' not found.", file_path)
                raise FileNotFoundError(f"Questions file '{file_path}' not found.")

            df = pd.read_excel(file_path)
            questions = df.iloc[:, 0].dropna().tolist()
            logging.info("Loaded %d questions from '%s'.", len(questions), file_path)
            return questions
        except Exception as e:
            logging.error(
                "Error loading questions from '%s': %s", file_path, e, exc_info=True
            )
            raise

    def load_transcript(self, file_path: str) -> str:
        """Load transcript from Word document"""
        from docx import Document

        try:
            if not os.path.exists(file_path):
                logging.error("Transcript file '%s' not found.", file_path)
                raise FileNotFoundError(f"Transcript file '{file_path}' not found.")

            document = Document(file_path)
            transcript = " ".join(
                para.text.strip() for para in document.paragraphs if para.text.strip()
            )
            logging.info("Loaded transcript from '%s'.", file_path)
            return transcript
        except Exception as e:
            logging.error(
                "Error loading transcript from '%s': %s", file_path, e, exc_info=True
            )
            raise

    def save_anonymized_transcript(
        self, anonymized_text: str, output_file: str
    ) -> None:
        """Save the anonymized transcript to a Word document"""
        from docx import Document

        try:
            document = Document()
            document.add_paragraph(anonymized_text)
            document.save(output_file)
            logging.info("Anonymized transcript saved to '%s'.", output_file)
        except Exception as e:
            logging.error(
                "Error saving anonymized transcript to '%s': %s",
                output_file,
                e,
                exc_info=True,
            )
            raise

    def _get_contextual_entities(self, text: str) -> List[Dict]:
        """Identify entities using context-aware methods combining multiple models"""
        try:
            entities = []

            # Get SpaCy entities
            doc = self.spacy_model(text)
            for ent in doc.ents:
                if ent.label_ in self.entity_type_map:
                    entities.append(
                        {
                            "text": ent.text,
                            "start": ent.start_char,
                            "end": ent.end_char,
                            "type": self.entity_type_map[ent.label_],
                            "confidence": (
                                ent._.probability
                                if hasattr(ent._, "probability")
                                else 0.8
                            ),
                            "source": "spacy",
                            "role": self.grammatical_roles.get(
                                ent.root.dep_, "UNKNOWN"
                            ),
                        }
                    )

            # Get transformer-based entities
            transformer_entities = self.ner_pipeline(text)
            for ent in transformer_entities:
                mapped_type = self.entity_type_map.get(ent["entity_group"], "MISC")
                entities.append(
                    {
                        "text": ent["word"],
                        "start": ent["start"],
                        "end": ent["end"],
                        "type": mapped_type,
                        "confidence": ent["score"],
                        "source": "transformer",
                    }
                )

            # Resolve overlapping entities and enhance with context
            resolved_entities = self._resolve_entity_conflicts(entities)
            enhanced_entities = self._enhance_with_context(text, resolved_entities)

            return enhanced_entities

        except Exception as e:
            logging.error("Error in contextual entity extraction: %s", e, exc_info=True)
            return []

    def _resolve_entity_conflicts(self, entities: List[Dict]) -> List[Dict]:
        """Resolve overlapping entity detections using confidence scores and model priority"""
        sorted_entities = sorted(entities, key=lambda x: (x["start"], -x["confidence"]))
        resolved = []
        last_end = -1

        for entity in sorted_entities:
            if entity["start"] >= last_end:
                resolved.append(entity)
                last_end = entity["end"]
            else:
                # Check for nested entities
                prev_entity = resolved[-1]
                if entity["confidence"] > prev_entity["confidence"]:
                    resolved[-1] = entity
                    last_end = entity["end"]

        return resolved

    def _determine_entity_type(
        self, token: "spacy.tokens.Token", doc: "spacy.tokens.Doc"
    ) -> Optional[str]:
        """Determine entity type based on token context and dependencies"""
        # Check if token is part of a named entity
        if token.ent_type_:
            return self.entity_type_map.get(token.ent_type_)

        # Check for role indicators
        role_indicators = {
            "manager",
            "director",
            "supervisor",
            "lead",
            "head",
            "coordinator",
            "specialist",
            "analyst",
            "engineer",
            "developer",
            "consultant",
            "administrator",
        }

        # Check for organization indicators
        org_indicators = {
            "company",
            "corp",
            "inc",
            "ltd",
            "llc",
            "corporation",
            "associates",
            "partners",
            "group",
            "solutions",
        }

        # Check the token and its children
        token_text = token.text.lower()

        # Direct role match
        if token_text in role_indicators:
            return "ROLE"

        # Direct organization match
        if token_text in org_indicators:
            return "ORGANIZATION"

        # Check compound structure
        compound_text = ""
        for child in token.children:
            if child.dep_ == "compound":
                compound_text += child.text.lower() + " "
        compound_text += token_text

        # Check compound text for indicators
        if any(indicator in compound_text for indicator in role_indicators):
            return "ROLE"
        if any(indicator in compound_text for indicator in org_indicators):
            return "ORGANIZATION"

        # Check parent token for context
        if token.dep_ == "compound" and token.head.ent_type_:
            return self.entity_type_map.get(token.head.ent_type_)

        return None

    def _enhance_with_context(self, text: str, entities: List[Dict]) -> List[Dict]:
        """Enhance entity detection using contextual clues, patterns, and dependency parsing"""
        doc = self.spacy_model(text)
        enhanced_entities = entities.copy()

        # Context patterns
        role_patterns = [
            r"(?i)(?:^|\s)(senior|junior|lead|chief|head\sof|director\sof|manager\sof|vp\sof)[\s\w]+",
            r"(?i)(?:^|\s)(\w+ist|\w+eer|\w+ctor)\b",
            r"(?i)(?:^|\s)(account|sales|marketing|support|customer|service)\s+(?:representative|manager|executive)\b",
        ]

        project_patterns = [
            r"(?i)(?:project|initiative|program)\s+([A-Z][A-Za-z0-9_\-]+)",
            r"(?i)([A-Z][A-Za-z0-9_\-]+)\s+(?:project|initiative|program)\b",
        ]

        # Process dependency tree for compound names and their context
        for token in doc:
            if token.dep_ == "compound":
                compound_phrase = []
                current = token
                while current.head != current and len(compound_phrase) < 5:
                    compound_phrase.append(current.text)
                    current = current.head
                compound_phrase.append(current.text)

                entity_text = " ".join(reversed(compound_phrase))
                entity_type = self._determine_entity_type(current, doc)

                if entity_type:
                    enhanced_entities.append(
                        {
                            "text": entity_text,
                            "start": token.idx,
                            "end": current.idx + len(current.text),
                            "type": entity_type,
                            "role": self.grammatical_roles.get(current.dep_, "UNKNOWN"),
                            "confidence": 0.9,
                            "source": "dependency",
                        }
                    )

            # Process named entities with their grammatical roles
            if token.ent_type_ and token.ent_type_ in self.entity_type_map:
                enhanced_entities.append(
                    {
                        "text": token.text,
                        "start": token.idx,
                        "end": token.idx + len(token.text),
                        "type": self.entity_type_map[token.ent_type_],
                        "role": self.grammatical_roles.get(token.dep_, "UNKNOWN"),
                        "confidence": 0.85,
                        "source": "grammatical",
                    }
                )

        # Process appositive phrases
        for token in doc:
            if token.dep_ == "appos" and token.head.ent_type_:
                enhanced_entities.append(
                    {
                        "text": f"{token.head.text} ({token.text})",
                        "start": token.head.idx,
                        "end": token.idx + len(token.text),
                        "type": self.entity_type_map.get(token.head.ent_type_, "MISC"),
                        "role": "APPOSITIVE",
                        "confidence": 0.9,
                        "source": "appositive",
                    }
                )

        # Apply pattern-based detection
        for pattern in role_patterns + project_patterns:
            for match in re.finditer(pattern, text):
                enhanced_entities.append(
                    {
                        "text": match.group(0).strip(),
                        "start": match.start(),
                        "end": match.end(),
                        "type": "ROLE" if pattern in role_patterns else "PROJECT",
                        "confidence": 0.85,
                        "source": "pattern",
                    }
                )

        return self._resolve_entity_conflicts(enhanced_entities)

    def anonymize_text(self, text: str, custom_words: Dict[str, str] = None) -> str:
        """Enhanced anonymization using multiple NLP models and contextual understanding"""
        try:
            # Get all entities using enhanced detection
            entities = self._get_contextual_entities(text)

            # Add custom word replacements
            if custom_words:
                for word, replacement in custom_words.items():
                    for match in re.finditer(
                        rf"\b{re.escape(word)}\b", text, re.IGNORECASE
                    ):
                        entities.append(
                            {
                                "text": match.group(0),
                                "start": match.start(),
                                "end": match.end(),
                                "type": "CUSTOM",
                                "replacement": replacement,
                                "confidence": 1.0,
                                "source": "custom",
                            }
                        )

            # Sort entities by start position (reversed for replacing from end)
            entities.sort(key=lambda x: -x["start"])

            # Apply replacements with role context when available
            anonymized = text
            for entity in entities:
                if "role" in entity and entity["role"] != "UNKNOWN":
                    replacement = (
                        entity["replacement"]
                        if "replacement" in entity
                        else f"[{entity['type']}-{entity['role']}]"
                    )
                else:
                    replacement = (
                        entity["replacement"]
                        if "replacement" in entity
                        else f"[{entity['type']}]"
                    )
                anonymized = (
                    anonymized[: entity["start"]]
                    + replacement
                    + anonymized[entity["end"] :]
                )

            # Log anonymization statistics with role information
            entity_counts = {}
            role_counts = {}
            for entity in entities:
                entity_type = entity["type"]
                entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1

                if "role" in entity and entity["role"] != "UNKNOWN":
                    role = entity["role"]
                    role_counts[role] = role_counts.get(role, 0) + 1

            # Log detailed statistics
            logging.info(
                "Anonymization complete: %s",
                ", ".join(f"{k}: {v}" for k, v in entity_counts.items()),
            )
            if role_counts:
                logging.info(
                    "Grammatical roles found: %s",
                    ", ".join(f"{k}: {v}" for k, v in role_counts.items()),
                )

            return anonymized

        except Exception as e:
            logging.error("Error in enhanced anonymization: %s", e, exc_info=True)
            return text

    def process_single_transcript(
        self,
        transcript_file: str,
        questions: List[str],
        anonymize: bool = True,
        progress_callback: Optional[Callable] = None,
        custom_words: Dict[str, str] = None,
        analyze_sentiment: bool = False,
        analyze_topics: bool = False,
    ) -> Dict[int, Dict[str, str]]:
        """Process transcript with enhanced emotion analysis"""
        try:
            if progress_callback:
                progress_callback("status", "Loading transcript...")

            # Load and process transcript
            transcript = self.load_transcript(transcript_file)

            # Do topic analysis first if enabled
            if analyze_topics:
                try:
                    if self.topic_analyzer is None:
                        self.init_topic_analyzer()
                    logging.info(f"Starting topic analysis for {transcript_file}")
                    topic_results = self.topic_analyzer.analyze_topics(
                        transcript, os.path.basename(transcript_file), progress_callback
                    )
                    if progress_callback and hasattr(
                        self.topic_analyzer, "results_cache"
                    ):
                        progress_callback(
                            "topics_update", self.topic_analyzer.results_cache
                        )

                    logging.info(f"Topic analysis completed for {transcript_file}")
                    if progress_callback:
                        progress_callback("status", "Analyzing topics...")
                        progress_callback("topics_loading", True)

                    logging.info(f"Starting topic analysis for {transcript_file}")
                    topic_results = self.topic_analyzer.analyze_topics(
                        transcript, os.path.basename(transcript_file), progress_callback
                    )
                    logging.info(f"Topic analysis completed for {transcript_file}")

                    if progress_callback:
                        progress_callback(
                            "topics_update", self.topic_analyzer.results_cache
                        )
                        progress_callback("topics_loading", False)
                except Exception as e:
                    logging.error(f"Topic analysis failed: {str(e)}", exc_info=True)
                    if progress_callback:
                        progress_callback("topics_loading", False)

            # Continue with regular processing
            entity_mapping = {}
            if anonymize:
                entities = self._get_contextual_entities(transcript)
                for entity in entities:
                    entity_type = entity["type"]
                    anon_text = f"[{entity_type}]"
                    entity_mapping[anon_text] = entity["text"]

                if custom_words:
                    for word, replacement in custom_words.items():
                        entity_mapping[replacement] = word

            # Anonymize transcript
            anonymized_transcript = self.anonymize_text(transcript, custom_words)

            if anonymize:
                anonymized_output_path = os.path.join(
                    os.path.dirname(transcript_file),
                    f"anonymized_{os.path.basename(transcript_file)}",
                )
                self.save_anonymized_transcript(
                    anonymized_transcript, anonymized_output_path
                )

            if progress_callback:
                progress_callback("status", "Starting question analysis...")

            # Get answers from GPT
            answers = self.get_answers_from_gpt(
                anonymized_transcript, questions, progress_callback
            )

            # Handle sentiment analysis if requested
            if analyze_sentiment:
                if progress_callback:
                    progress_callback("status", "Analyzing emotional content...")

                for idx, answer_data in answers.items():
                    try:
                        answer_text = answer_data["answer"]
                        deanonymized_answer = self._deanonymize_text(
                            answer_text, entity_mapping
                        )
                        emotion_analysis = self.analyze_answer_sentiment(
                            deanonymized_answer
                        )

                        if (
                            isinstance(emotion_analysis, dict)
                            and "sentiment_summary" in emotion_analysis
                        ):
                            answers[idx]["sentiment_summary"] = emotion_analysis[
                                "sentiment_summary"
                            ]
                        else:
                            answers[idx][
                                "sentiment_summary"
                            ] = "Error in sentiment analysis"

                    except Exception as e:
                        logging.error(
                            f"Error processing sentiment for answer {idx}: {str(e)}",
                            exc_info=True,
                        )
                        answers[idx][
                            "sentiment_summary"
                        ] = f"Error in sentiment analysis: {str(e)}"

            return answers

        except Exception as e:
            logging.error(
                f"Error processing transcript '{transcript_file}': {str(e)}",
                exc_info=True,
            )
            raise

    def _deanonymize_text(self, text: str, entity_mapping: Dict[str, str]) -> str:
        """Helper method to deanonymize text using entity mapping"""
        deanonymized = text
        quote_pattern = r'"([^"]*)"'

        def replace_in_quote(match):
            quote = match.group(1)
            replaced_quote = quote
            for anon_text, original_text in entity_mapping.items():
                replaced_quote = replaced_quote.replace(anon_text, original_text)
            return f'"{replaced_quote}"'

        return re.sub(quote_pattern, replace_in_quote, deanonymized)

    def get_answers_from_gpt(
        self,
        transcript: str,
        questions: List[str],
        progress_callback: Optional[Callable] = None,
    ) -> Dict[int, Dict[str, str]]:
        """Get answers using GPT with enhanced system prompt and progress tracking"""
        if not self.client:
            logging.error("OpenAI client not initialized.")
            raise ValueError("OpenAI client not initialized. Please set API key first.")

        answers = {}
        total_questions = len(questions)
        available_tokens = self.max_tokens - 1000
        tokens_per_question = max(2000, available_tokens // total_questions)

        for i, question in enumerate(questions, 1):
            try:
                if progress_callback:
                    progress_callback(
                        "question_progress",
                        f"Processing question {i}/{total_questions}: {question[:50]}...",
                    )

                # Create enhanced prompt with anonymization context
                prompt = self._create_enhanced_prompt(question, transcript)

                # Initialize retry counter for this specific question
                question_retries = 0
                max_question_retries = 3
                success = False

                while not success and question_retries < max_question_retries:
                    try:
                        self.rate_limiter.wait()

                        response = self.client.chat.completions.create(
                            model=self.model_name,
                            messages=[
                                {
                                    "role": "system",
                                    "content": "You are an expert interview analyst skilled in extracting precise information and identifying timestamps. Always provide complete responses that include extensive direct quotes from the source material. Each key point should be supported by relevant quotes.",
                                },
                                {"role": "user", "content": prompt},
                            ],
                            temperature=0.3,
                            max_tokens=tokens_per_question,
                            presence_penalty=0.0,
                            frequency_penalty=0.0,
                        )

                        answer_text = response.choices[0].message.content.strip()

                        if not self._validate_answer(answer_text):
                            self.rate_limiter.wait()
                            prompt += "\n\nPlease ensure your response contains direct quotes and ends with a complete thought."
                            response = self.client.chat.completions.create(
                                model=self.model_name,
                                messages=[
                                    {
                                        "role": "system",
                                        "content": "You are an expert interview analyst. Provide complete responses with extensive direct quotes. Ensure every response includes multiple relevant quotes from the transcript.",
                                    },
                                    {"role": "user", "content": prompt},
                                ],
                                temperature=0.3,
                                max_tokens=tokens_per_question,
                                presence_penalty=0.0,
                                frequency_penalty=0.0,
                            )
                            answer_text = response.choices[0].message.content.strip()

                        success = True

                    except Exception as e:
                        question_retries += 1
                        if self.handle_rate_limit(e):
                            logging.warning(
                                f"Retrying question {i} after rate limit (Attempt {question_retries})"
                            )
                            continue
                        else:
                            raise

                if not success:
                    raise Exception(
                        f"Failed to process question after {max_question_retries} attempts"
                    )

                timestamp = self._extract_timestamp(answer_text)
                answers[i - 1] = {"answer": answer_text, "timestamp": timestamp}

                # Reset retry counter after successful processing
                self.retry_count = 0

            except Exception as e:
                logging.error(f"Error processing question {i}: {e}", exc_info=True)
                answers[i - 1] = {
                    "answer": f"Error processing question: {str(e)}",
                    "timestamp": "N/A",
                }

        return answers

    def _create_enhanced_prompt(self, question: str, transcript: str) -> str:
        """Create an enhanced prompt that includes context about anonymized entities"""
        return f"""Analyze this interview transcript and extract a precise answer for this question:

Question: {question}

Note: This transcript has been anonymized. Entity markers like [PERSON], [ORGANIZATION], etc., represent real names and sensitive information.

Transcript:
{transcript}

Instructions for Answer Extraction:
1. Identify all relevant portions of the transcript
2. Look for both direct and implied answers
3. Consider context from the entire conversation
4. Note the timestamp where the primary answer begins
5. Include any follow-up clarifications or additional details
6. Ensure response is complete and ends with a proper conclusion
7. When quoting anonymized sections, use the exact anonymized markers

Answer Format Requirements:
- Begin with the timestamp in [MM:SS] format
- Provide a clear, complete answer
- Use extensive direct quotes from the transcript, marking them with quotation marks (" ")
- When using quotes, include the speaker's exact words and context
- For longer quotes, use complete sentences or paragraphs to maintain context
- After each quote, explain its relevance to the question
- Maintain the original speaker's tone and intent
- Include contextual information when relevant
- End with a clear concluding statement

Please provide your response following this format:
[Timestamp] Complete answer with relevant quotes..."""

    def _validate_answer(self, answer: str) -> bool:
        """Validate the completeness and format of an answer"""
        return (
            answer.endswith((".", "!", "?", '"'))
            and '"' in answer
            and re.search(r"\[\d{1,2}:\d{2}\]", answer)
        )

    def _extract_timestamp(self, answer: str) -> str:
        """Extract the primary timestamp from an answer"""
        timestamp_pattern = r"\[?(\d{1,2}:\d{2})\]?"
        timestamp_matches = re.finditer(timestamp_pattern, answer)
        timestamps = [m.group(1) for m in timestamp_matches]
        return timestamps[0] if timestamps else "N/A"

    def _apply_formatting(self, worksheet) -> None:
        """Apply formatting to Excel worksheet"""
        try:
            from openpyxl.styles import PatternFill, Font, Alignment
            from openpyxl.utils import get_column_letter

            # Format headers
            header_fill = PatternFill(
                start_color="366092", end_color="366092", fill_type="solid"
            )
            header_font = Font(bold=True, color="FFFFFF")

            # Apply header formatting
            for cell in worksheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(wrap_text=True, vertical="center")

            # Set column widths
            worksheet.column_dimensions["A"].width = 40  # Question column
            worksheet.column_dimensions["B"].width = 15  # Timestamp column
            worksheet.column_dimensions["C"].width = 100  # Answer column
            worksheet.column_dimensions["D"].width = 50  # Summary column

            # Adjust row heights and wrap text for content
            for row in worksheet.iter_rows(min_row=2):
                worksheet.row_dimensions[row[0].row].height = 75
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True, vertical="top")

            logging.info("Excel formatting applied successfully.")
        except Exception as e:
            logging.error("Error applying Excel formatting: %s", e, exc_info=True)
            # Continue without formatting if there's an error
            pass

    def analyze_sentiment(self, text: str, chunk_size: int = 512) -> List[Dict]:
        """Analyze sentiment of text in chunks with windowing"""
        try:
            # Split text into sentences
            doc = self.spacy_model(text)
            sentences = [str(sent) for sent in doc.sents]

            # Process sentences with overlapping windows
            sentiments = []
            current_chunk = ""
            current_length = 0

            for sentence in sentences:
                # Add new sentence
                if current_length + len(sentence) <= chunk_size:
                    current_chunk += sentence + " "
                    current_length += len(sentence) + 1
                else:
                    # Process current chunk
                    if current_chunk:
                        result = self.sentiment_pipeline(current_chunk)[0]
                        sentiments.append(
                            {
                                "text": current_chunk.strip(),
                                "sentiment": result["label"],
                                "confidence": result["score"],
                                "timestamp": self._extract_timestamp(current_chunk)
                                or "N/A",
                            }
                        )

                    # Start new chunk
                    current_chunk = sentence + " "
                    current_length = len(sentence) + 1

            # Process final chunk
            if current_chunk:
                result = self.sentiment_pipeline(current_chunk)[0]
                sentiments.append(
                    {
                        "text": current_chunk.strip(),
                        "sentiment": result["label"],
                        "confidence": result["score"],
                        "timestamp": self._extract_timestamp(current_chunk) or "N/A",
                    }
                )

            return sentiments

        except Exception as e:
            logging.error("Error in sentiment analysis: %s", e, exc_info=True)
            raise

    def _generate_summary(self, answer: str) -> str:
        """Generate a brief summary of the answer"""
        try:
            self.rate_limiter.wait()

            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {
                        "role": "system",
                        "content": "Summarize the following in one or two sentences, maintaining any key quotes:",
                    },
                    {"role": "user", "content": answer},
                ],
                temperature=0.3,
                max_tokens=100,
            )
            summary = response.choices[0].message.content.strip()
            logging.debug("Generated summary for answer.")
            return summary
        except Exception as e:
            logging.error("Error generating summary: %s", e, exc_info=True)
            return "Summary unavailable"

    def save_results(
        self,
        answers: Dict[int, Dict[str, str]],
        questions: List[str],
        output_file: str,
        transcript_name: str,
    ) -> None:
        """Save results with sentiment summary in the same sheet"""
        import pandas as pd
        import os

        try:
            # Prepare data with sentiment
            data = []
            for idx, answer_data in answers.items():
                data_entry = {
                    "Question": questions[idx],
                    "Timestamp": answer_data["timestamp"],
                    "Answer": answer_data["answer"],
                    "Summary": self._generate_summary(answer_data["answer"]),
                }

                # Add sentiment summary if available
                if "sentiment_summary" in answer_data:
                    data_entry["Emotional Analysis"] = answer_data["sentiment_summary"]

                data.append(data_entry)

            df = pd.DataFrame(data)

            # Get sheet name from transcript filename
            sheet_name = os.path.splitext(os.path.basename(transcript_name))[0]
            sheet_name = sheet_name[:31]

            # Save to Excel
            try:
                with pd.ExcelWriter(
                    output_file, engine="openpyxl", mode="a", if_sheet_exists="replace"
                ) as writer:
                    df.to_excel(writer, index=False, sheet_name=sheet_name)
                    self._apply_formatting(writer.sheets[sheet_name])
            except FileNotFoundError:
                with pd.ExcelWriter(output_file, engine="openpyxl") as writer:
                    df.to_excel(writer, index=False, sheet_name=sheet_name)
                    self._apply_formatting(writer.sheets[sheet_name])

            logging.info("Results saved to '%s'", output_file)

        except Exception as e:
            logging.error("Error saving results: %s", e, exc_info=True)
            raise


class TimeEstimator:
    """Advanced time estimator with ML-inspired prediction and adaptive calibration"""

    def __init__(self):
        self.start_time = None
        self.processed_items = 0
        self.total_items = 0
        self.processing_history = []
        self.max_history_size = 20  # Increased for better trend analysis
        self.weight_decay = 0.92  # Adjusted for smoother decay
        self.last_update_time = None
        self.initial_estimate = None
        self.last_progress = 0
        self.confidence_threshold = 0.7
        self.min_samples_for_prediction = 3
        self.trend_window = 5
        self.kalman_gain = 0.3
        self.previous_estimate = None
        self.estimate_history = []
        self.max_estimate_history = 10
        self.outlier_threshold = 2.0  # Standard deviations

    def start(self, total_items):
        """Start timing a new batch of items with enhanced initialization"""
        self.start_time = time.time()
        self.last_update_time = self.start_time
        self.total_items = total_items
        self.processed_items = 0
        self.processing_history = []
        self.initial_estimate = None
        self.last_progress = 0
        self.previous_estimate = None
        self.estimate_history = []

    def _detect_outliers(self, speeds):
        """Detect and remove outliers using z-score method"""
        if len(speeds) < 4:  # Need enough samples for meaningful statistics
            return speeds

        mean = sum(speeds) / len(speeds)
        std = (sum((x - mean) ** 2 for x in speeds) / len(speeds)) ** 0.5

        return [s for s in speeds if abs((s - mean) / std) < self.outlier_threshold]

    def _calculate_trend(self):
        """Calculate processing speed trend using linear regression"""
        if len(self.processing_history) < self.trend_window:
            return 0

        recent_history = self.processing_history[-self.trend_window :]
        x = list(range(len(recent_history)))
        y = [entry["items_per_second"] for entry in recent_history]

        # Simple linear regression
        n = len(x)
        sum_x = sum(x)
        sum_y = sum(y)
        sum_xy = sum(i * j for i, j in zip(x, y))
        sum_xx = sum(i * i for i in x)

        try:
            slope = (n * sum_xy - sum_x * sum_y) / (n * sum_xx - sum_x * sum_x)
            return slope
        except ZeroDivisionError:
            return 0

    def _calculate_confidence(self, speeds):
        """Calculate confidence score based on consistency of measurements"""
        if len(speeds) < 2:
            return 0.5

        mean = sum(speeds) / len(speeds)
        variance = sum((s - mean) ** 2 for s in speeds) / len(speeds)
        cv = (variance**0.5) / mean if mean != 0 else float("inf")

        # Convert coefficient of variation to confidence score
        confidence = max(0, min(1, 1 - (cv / 2)))
        return confidence

    def _apply_kalman_filter(self, new_estimate):
        """Apply Kalman filter to smooth estimates"""
        if self.previous_estimate is None:
            self.previous_estimate = new_estimate
            return new_estimate

        kalman_estimate = self.previous_estimate + self.kalman_gain * (
            new_estimate - self.previous_estimate
        )
        self.previous_estimate = kalman_estimate
        return kalman_estimate

    def update(self, items_completed):
        """Update progress with enhanced statistical processing"""
        if self.start_time is None or items_completed < self.last_progress:
            return None

        current_time = time.time()
        newly_processed = items_completed - self.last_progress

        if newly_processed > 0:
            duration = max(0.001, current_time - self.last_update_time)

            # Calculate speed and detect if it's an outlier
            speed = newly_processed / duration
            if self.processing_history:
                speeds = [
                    entry["items_per_second"] for entry in self.processing_history
                ]
                mean_speed = sum(speeds) / len(speeds)
                std_speed = (
                    sum((s - mean_speed) ** 2 for s in speeds) / len(speeds)
                ) ** 0.5

                # Only add to history if not an extreme outlier
                if abs(speed - mean_speed) <= self.outlier_threshold * std_speed:
                    self.processing_history.append(
                        {
                            "timestamp": current_time,
                            "items_processed": newly_processed,
                            "duration": duration,
                            "items_per_second": speed,
                            "cumulative_progress": items_completed,
                        }
                    )
            else:
                # First entry
                self.processing_history.append(
                    {
                        "timestamp": current_time,
                        "items_processed": newly_processed,
                        "duration": duration,
                        "items_per_second": speed,
                        "cumulative_progress": items_completed,
                    }
                )

            # Maintain history size
            if len(self.processing_history) > self.max_history_size:
                self.processing_history.pop(0)

            self.last_update_time = current_time
            self.processed_items = items_completed
            self.last_progress = items_completed

    def get_estimate(self, current_progress):
        """Get estimated time remaining using advanced statistical methods"""
        if not self.processing_history or current_progress >= 100:
            return None

        try:
            if current_progress < 0 or current_progress > 100:
                return None

            # Get speed measurements and remove outliers
            speeds = [entry["items_per_second"] for entry in self.processing_history]
            cleaned_speeds = self._detect_outliers(speeds)

            if not cleaned_speeds:
                return None

            # Calculate weighted average speed with exponential decay
            total_weight = 0
            weighted_speed = 0

            for idx, speed in enumerate(cleaned_speeds):
                weight = self.weight_decay ** (len(cleaned_speeds) - idx - 1)
                weighted_speed += speed * weight
                total_weight += weight

            avg_speed = weighted_speed / total_weight if total_weight > 0 else 0

            if avg_speed <= 0:
                return None

            # Calculate trend and adjust speed
            trend = self._calculate_trend()
            adjusted_speed = max(0.001, avg_speed + (trend * len(cleaned_speeds) / 2))

            # Calculate remaining items and basic estimate
            remaining_progress = 100 - current_progress
            remaining_items = (remaining_progress / 100) * self.total_items
            base_estimate = remaining_items / adjusted_speed

            # Calculate confidence score
            confidence = self._calculate_confidence(cleaned_speeds)

            # Apply adaptive blending based on confidence
            if self.initial_estimate is None:
                self.initial_estimate = base_estimate
                estimated_seconds = base_estimate
            else:
                completion_factor = current_progress / 100
                confidence_weight = confidence * completion_factor

                estimated_seconds = (
                    base_estimate * confidence_weight
                    + self.initial_estimate * (1 - confidence_weight)
                )

            # Apply Kalman filtering
            estimated_seconds = self._apply_kalman_filter(estimated_seconds)

            # Store estimate history for trend analysis
            self.estimate_history.append(estimated_seconds)
            if len(self.estimate_history) > self.max_estimate_history:
                self.estimate_history.pop(0)

            # Calculate median of recent estimates for stability
            if len(self.estimate_history) >= 3:
                sorted_estimates = sorted(self.estimate_history)
                median_estimate = sorted_estimates[len(sorted_estimates) // 2]
                # Blend current estimate with median for stability
                estimated_seconds = estimated_seconds * 0.7 + median_estimate * 0.3

            # Add safeguards
            estimated_seconds = max(0, min(estimated_seconds, 86400 * 365))

            # Format and return the estimate
            return self.format_time_estimate(estimated_seconds)

        except Exception as e:
            logging.error(f"Error calculating time estimate: {e}")
            return None

    @staticmethod
    def format_time_estimate(seconds):
        """Format time estimate with improved granularity and natural language"""
        if seconds < 30:
            return "Less than a minute"
        elif seconds < 60:
            return "About a minute"
        elif seconds < 3600:
            minutes = seconds / 60
            if minutes < 5:
                return f"About {max(1, int(minutes))} minutes"
            else:
                rounded_minutes = int(minutes / 5) * 5
                return f"About {rounded_minutes} minutes"
        elif seconds < 7200:  # Less than 2 hours
            hours = int(seconds / 3600)
            minutes = int((seconds % 3600) / 60)
            rounded_minutes = int(minutes / 5) * 5
            if rounded_minutes == 60:
                hours += 1
                rounded_minutes = 0
            if rounded_minutes == 0:
                return f"About {hours} hour"
            return f"About {hours} hour {rounded_minutes} minutes"
        elif seconds < 86400:  # Less than 24 hours
            hours = int(seconds / 3600)
            return f"About {hours} hours"
        else:
            days = int(seconds / 86400)
            hours = int((seconds % 86400) / 3600)
            if hours >= 12:
                days += 1
                return f"About {days} days"
            elif hours > 0:
                return f"About {days} days {hours} hours"
            return f"About {days} days"

    def get_processing_statistics(self):
        """Get detailed processing statistics with confidence metrics"""
        if not self.processing_history:
            return None

        speeds = [entry["items_per_second"] for entry in self.processing_history]
        cleaned_speeds = self._detect_outliers(speeds)

        stats = {
            "avg_speed": (
                sum(cleaned_speeds) / len(cleaned_speeds) if cleaned_speeds else 0
            ),
            "min_speed": min(cleaned_speeds) if cleaned_speeds else 0,
            "max_speed": max(cleaned_speeds) if cleaned_speeds else 0,
            "total_processed": self.processed_items,
            "elapsed_time": time.time() - self.start_time,
            "confidence_score": self._calculate_confidence(cleaned_speeds),
            "trend": self._calculate_trend(),
            "samples_collected": len(self.processing_history),
            "outliers_removed": len(speeds) - len(cleaned_speeds),
        }

        if len(cleaned_speeds) >= 2:
            mean = stats["avg_speed"]
            variance = sum((s - mean) ** 2 for s in cleaned_speeds) / len(
                cleaned_speeds
            )
            stats["speed_variance"] = variance
            stats["coefficient_of_variation"] = (
                (variance**0.5) / mean if mean != 0 else float("inf")
            )

        return stats


def setup_topic_modeling_dependencies() -> None:
    """Install additional dependencies required for topic modeling"""
    import subprocess
    import sys
    import logging
    from tkinter import messagebox

    topic_packages = [
        "bertopic>=0.15.0",
        "sentence-transformers>=2.2.2",
        "umap-learn>=0.5.3",
        "hdbscan>=0.8.29",
        "plotly>=5.13.0",
    ]

    for package in topic_packages:
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", "--no-cache-dir", package]
            )
            logging.info(f"Successfully installed {package}")
        except Exception as e:
            error_msg = f"Failed to install {package}: {str(e)}"
            logging.error(error_msg)
            messagebox.showerror("Installation Error", error_msg)
            raise


class TopicAnalyzer:
    """Advanced topic modeling using BERTopic with local processing"""

    def __init__(self):
        """Initialize topic analyzer with required models and settings"""
        try:
            from bertopic import BERTopic
            from sentence_transformers import SentenceTransformer
            import torch
            import numpy as np
            import logging

            # Initialize the BERT encoder
            self.device = "cuda" if torch.cuda.is_available() else "cpu"
            logging.info(f"Using device: {self.device}")

            self.encoder = SentenceTransformer("all-MiniLM-L6-v2", device=self.device)
            logging.info("Initialized SentenceTransformer")

            # Initialize BERTopic with optimized parameters
            self.topic_model = BERTopic(
                embedding_model=self.encoder,
                min_topic_size=5,
                n_gram_range=(1, 3),
                calculate_probabilities=True,
                verbose=True,
            )
            logging.info("Initialized BERTopic model")

            self.results_cache = {}

        except Exception as e:
            logging.error(f"Error initializing TopicAnalyzer: {str(e)}", exc_info=True)
            raise

    def preprocess_transcript(self, text: str) -> List[str]:
        """Split transcript into meaningful chunks for topic modeling"""
        try:
            import spacy
            import logging

            # Load spaCy model for better sentence segmentation
            nlp = spacy.load("en_core_web_lg")
            doc = nlp(text)

            # Create chunks based on natural language boundaries
            chunks = []
            current_chunk = []
            current_length = 0
            target_chunk_size = 200  # Optimal chunk size for BERT

            for sent in doc.sents:
                sent_text = sent.text.strip()
                sent_length = len(sent_text.split())

                if current_length + sent_length > target_chunk_size:
                    if current_chunk:
                        chunks.append(" ".join(current_chunk))
                    current_chunk = [sent_text]
                    current_length = sent_length
                else:
                    current_chunk.append(sent_text)
                    current_length += sent_length

            # Add the last chunk if it exists
            if current_chunk:
                chunks.append(" ".join(current_chunk))

            logging.info(f"Created {len(chunks)} chunks from transcript")
            return chunks

        except Exception as e:
            logging.error(f"Error in transcript preprocessing: {str(e)}", exc_info=True)
            raise

    def analyze_topics(
        self,
        transcript: str,
        transcript_name: str,
        progress_callback: Optional[Callable] = None,
    ) -> Dict[str, Any]:
        """Perform advanced topic modeling on a transcript"""
        import numpy as np

        try:
            if progress_callback:
                progress_callback(
                    "status", "Preprocessing transcript for topic analysis..."
                )
                progress_callback("topics_loading", True)

            logging.info(f"Starting topic analysis for {transcript_name}")

            # Preprocess and chunk the transcript
            chunks = self.preprocess_transcript(transcript)

            if progress_callback:
                progress_callback("status", "Generating embeddings...")

            # Generate embeddings
            embeddings = self.encoder.encode(
                chunks, show_progress_bar=False, device=self.device
            )

            if progress_callback:
                progress_callback("status", "Identifying topics...")

            # Fit the topic model and ensure outputs are numpy arrays
            topics, probs = self.topic_model.fit_transform(chunks, embeddings)
            topics = np.array(topics)
            probs = np.array(probs)

            # Get topic insights
            topic_info = self.topic_model.get_topic_info()
            topic_repr = self.topic_model.get_topics()

            # Calculate topic distribution
            topic_sizes = topic_info["Count"].to_dict()
            total_docs = sum(topic_sizes.values())
            topic_distribution = {
                topic_id: count / total_docs
                for topic_id, count in topic_sizes.items()
                if topic_id != -1  # Exclude noise topic
            }

            # Process results
            results = {
                "transcript_name": transcript_name,
                "topics": topics,
                "probabilities": probs,
                "topic_info": topic_info.to_dict(),
                "chunks": chunks,
                "topic_representatives": {
                    str(topic_id): [word for word, _ in words[:10]]
                    for topic_id, words in topic_repr.items()
                    if topic_id != -1  # Exclude noise topic
                },
                "topic_distribution": topic_distribution,
            }

            # Calculate topic coherence and diversity metrics
            results["metrics"] = {
                "coherence": self.calculate_coherence(results),
                "diversity": self.calculate_diversity(results),
            }

            # Store results in cache and update GUI
            self.results_cache[transcript_name] = results

            if progress_callback:
                logging.info(f"Sending topics update for {transcript_name}")
                progress_callback("topics_update", self.results_cache)
                progress_callback("topics_loading", False)

            logging.info(f"Topic analysis completed for {transcript_name}")
            return results

        except Exception as e:
            logging.error(f"Error in topic analysis: {str(e)}", exc_info=True)
            if progress_callback:
                progress_callback("topics_loading", False)
            raise

    def calculate_coherence(self, results: Dict) -> float:
        """Calculate topic coherence using normalized pointwise mutual information"""
        try:
            import numpy as np
            import logging

            coherence_scores = []

            for topic_id, words in results["topic_representatives"].items():
                try:
                    topic_coherence = self.topic_model.compute_coherence_score(words)
                    if np.isfinite(topic_coherence):  # Check for valid numeric value
                        coherence_scores.append(topic_coherence)
                except Exception as e:
                    logging.warning(
                        f"Error calculating coherence for topic {topic_id}: {str(e)}"
                    )
                    continue

            if coherence_scores:
                return float(np.mean(coherence_scores))
            return 0.0

        except Exception as e:
            logging.error(f"Error calculating coherence: {str(e)}")
            return 0.0

    def calculate_diversity(self, results: Dict) -> float:
        """Calculate topic diversity using unique words ratio"""
        try:
            all_words = set()
            total_words = 0

            for words in results["topic_representatives"].values():
                # Convert words to lowercase for better comparison
                words_lower = [word.lower() for word in words]
                all_words.update(words_lower)
                total_words += len(words)

            return len(all_words) / total_words if total_words > 0 else 0.0

        except Exception as e:
            logging.error(f"Error calculating diversity: {str(e)}")
            return 0.0

    def get_representative_chunks(
        self, topic_id: str, results: Dict[str, Any]
    ) -> List[str]:
        """Get representative text chunks for a topic"""
        try:
            # Convert topic_id to integer for indexing
            topic_id_int = int(topic_id)

            # Convert topics to numpy array if it isn't already
            topics_array = np.array(results["topics"])

            # Find indices where topics match
            topic_idx = np.where(topics_array == topic_id_int)[0]

            if len(topic_idx) == 0:
                return []

            # Ensure probabilities is a numpy array
            probs_array = np.array(results["probabilities"])

            # Get probabilities for this topic
            chunk_probs = probs_array[topic_idx]

            # Sort indices by probability
            sorted_indices = topic_idx[np.argsort(chunk_probs)[-3:]]

            # Convert to Python int for indexing
            sorted_indices = [int(idx) for idx in sorted_indices]

            # Get the chunks
            return [results["chunks"][idx] for idx in sorted_indices]

        except (ValueError, TypeError, IndexError) as e:
            logging.error(f"Error getting representative chunks: {str(e)}")
            return []
        except Exception as e:
            logging.error(f"Unexpected error in get_representative_chunks: {str(e)}")
            return []

    def clear_cache(self, transcript_name: Optional[str] = None) -> None:
        """Clear the results cache for a specific transcript or all transcripts"""
        try:
            if transcript_name is not None:
                if transcript_name in self.results_cache:
                    del self.results_cache[transcript_name]
                    logging.info(f"Cleared cache for transcript: {transcript_name}")
            else:
                self.results_cache.clear()
                logging.info("Cleared entire results cache")
        except Exception as e:
            logging.error(f"Error clearing cache: {str(e)}")
            raise


class TopicsTab:
    def __init__(self, notebook, queue):
        self.topics_frame = ttk.Frame(notebook, padding=10)
        notebook.add(self.topics_frame, text="Topics")
        self.root = notebook.winfo_toplevel()
        self.setup_topics_tab()
        self.queue = queue
        self.current_results = {}
        self.loading = False

    def setup_topics_tab(self):
        self.topics_frame.columnconfigure(0, weight=1)
        self.topics_frame.rowconfigure(1, weight=1)

        # Transcript selector
        control_frame = ttk.Frame(self.topics_frame)
        control_frame.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        control_frame.columnconfigure(1, weight=1)

        ttk.Label(control_frame, text="Select Transcript:").grid(
            row=0, column=0, sticky="w"
        )
        self.transcript_var = tk.StringVar()
        self.transcript_combo = ttk.Combobox(
            control_frame, textvariable=self.transcript_var, state="readonly"
        )
        self.transcript_combo.grid(row=0, column=1, sticky="ew", padx=5)
        self.transcript_combo.bind("<<ComboboxSelected>>", self.on_transcript_selected)

        # Main content
        content_frame = ttk.Frame(self.topics_frame)
        content_frame.grid(row=1, column=0, sticky="nsew")
        content_frame.columnconfigure(1, weight=1)
        content_frame.rowconfigure(0, weight=1)

        # Topics list with scrollbar
        topics_list_frame = ttk.LabelFrame(content_frame, text="Topics", padding=5)
        topics_list_frame.grid(row=0, column=0, sticky="ns", padx=(0, 5))

        topics_scroll = ttk.Scrollbar(topics_list_frame)
        topics_scroll.pack(side="right", fill="y")

        self.topics_listbox = tk.Listbox(
            topics_list_frame,
            width=30,
            selectmode=tk.SINGLE,
            exportselection=False,
            yscrollcommand=topics_scroll.set,
        )
        self.topics_listbox.pack(side="left", fill="both", expand=True)
        topics_scroll.config(command=self.topics_listbox.yview)
        self.topics_listbox.bind("<<ListboxSelect>>", self.on_topic_selected)

        # Topic details section
        details_frame = ttk.Frame(content_frame)
        details_frame.grid(row=0, column=1, sticky="nsew")
        details_frame.columnconfigure(0, weight=1)
        details_frame.rowconfigure(1, weight=1)

        # Topic info section
        info_frame = ttk.LabelFrame(details_frame, text="Topic Information", padding=5)
        info_frame.grid(row=0, column=0, sticky="ew", pady=(0, 5))
        info_frame.columnconfigure(1, weight=1)

        self.topic_name_var = tk.StringVar()
        self.topic_name_label = ttk.Label(
            info_frame, textvariable=self.topic_name_var, font=("Helvetica", 12, "bold")
        )
        self.topic_name_label.grid(
            row=0, column=0, columnspan=2, sticky="w", pady=(0, 5)
        )

        self.topic_size_var = tk.StringVar()
        ttk.Label(info_frame, text="Size:").grid(row=1, column=0, sticky="w")
        ttk.Label(info_frame, textvariable=self.topic_size_var).grid(
            row=1, column=1, sticky="w"
        )

        self.topic_coherence_var = tk.StringVar()
        ttk.Label(info_frame, text="Coherence:").grid(row=2, column=0, sticky="w")
        ttk.Label(info_frame, textvariable=self.topic_coherence_var).grid(
            row=2, column=1, sticky="w"
        )

        # Keywords section with scrollbar
        keywords_frame = ttk.LabelFrame(details_frame, text="Keywords", padding=5)
        keywords_frame.grid(row=1, column=0, sticky="nsew", pady=(0, 5))
        keywords_frame.columnconfigure(0, weight=1)
        keywords_frame.rowconfigure(0, weight=1)

        keywords_scroll = ttk.Scrollbar(keywords_frame)
        keywords_scroll.grid(row=0, column=1, sticky="ns")

        self.keywords_text = tk.Text(
            keywords_frame, wrap=tk.WORD, height=4, yscrollcommand=keywords_scroll.set
        )
        self.keywords_text.grid(row=0, column=0, sticky="nsew")
        keywords_scroll.config(command=self.keywords_text.yview)
        self.keywords_text.configure(state="disabled")

        # Representative chunks section
        chunks_frame = ttk.LabelFrame(
            details_frame, text="Representative Text Chunks", padding=5
        )
        chunks_frame.grid(row=2, column=0, sticky="nsew")
        chunks_frame.columnconfigure(0, weight=1)
        chunks_frame.rowconfigure(0, weight=1)

        self.chunks_text = scrolledtext.ScrolledText(
            chunks_frame, wrap=tk.WORD, height=10
        )
        self.chunks_text.pack(fill="both", expand=True)
        self.chunks_text.configure(state="disabled")

    def update_transcript_list(self, results_cache):
        """Update transcript dropdown with available results"""
        if not results_cache:
            self.transcript_combo["values"] = []
            self.transcript_var.set("")
            return

        transcripts = list(results_cache.keys())
        self.transcript_combo["values"] = transcripts

        if transcripts and not self.transcript_var.get():
            self.transcript_var.set(transcripts[0])
            self.on_transcript_selected(None)

    def on_transcript_selected(self, event):
        """Handle transcript selection"""
        transcript = self.transcript_var.get()
        if not transcript or transcript not in self.current_results:
            return

        results = self.current_results[transcript]
        self.topics_listbox.delete(0, tk.END)

        topic_reps = results.get("topic_representatives", {})
        sorted_topics = sorted(
            topic_reps.keys(), key=lambda x: int(x) if x.isdigit() else float("inf")
        )

        for topic_id in sorted_topics:
            words = topic_reps[topic_id]
            topic_name = f"Topic {topic_id}: {', '.join(words[:3])}..."
            self.topics_listbox.insert(tk.END, topic_name)

        if self.topics_listbox.size() > 0:
            self.topics_listbox.selection_set(0)
            self.on_topic_selected(None)

    def on_topic_selected(self, event):
        """Handle topic selection"""
        if self.loading:
            return

        selection = self.topics_listbox.curselection()
        if not selection:
            return

        transcript = self.transcript_var.get()
        if not transcript or transcript not in self.current_results:
            return

        results = self.current_results[transcript]
        sorted_topics = sorted(
            results["topic_representatives"].keys(),
            key=lambda x: int(x) if x.isdigit() else float("inf"),
        )
        topic_id = sorted_topics[selection[0]]

        # Update topic information
        self.topic_name_var.set(f"Topic {topic_id}")

        # Update size information
        distribution = results.get("topic_distribution", {}).get(topic_id, 0)
        self.topic_size_var.set(f"{distribution:.1%} of transcript")

        # Update coherence
        coherence = results.get("metrics", {}).get("coherence", 0)
        self.topic_coherence_var.set(f"{coherence:.3f}")

        # Update keywords
        self.keywords_text.configure(state="normal")
        self.keywords_text.delete(1.0, tk.END)
        keywords = results["topic_representatives"][topic_id]
        self.keywords_text.insert(1.0, ", ".join(keywords))
        self.keywords_text.configure(state="disabled")

        # Update representative chunks
        self.chunks_text.configure(state="normal")
        self.chunks_text.delete(1.0, tk.END)

        topic_chunks = []
        for i, topic in enumerate(results["topics"]):
            if str(topic) == str(topic_id):
                topic_chunks.append(results["chunks"][i])
                if len(topic_chunks) >= 3:
                    break

        for i, chunk in enumerate(topic_chunks, 1):
            self.chunks_text.insert(tk.END, f"Chunk {i}:\n{chunk}\n\n")
        self.chunks_text.configure(state="disabled")

    def update_results(self, results_cache):
        try:
            if not results_cache:
                logging.warning("Empty results cache received")
                return

            logging.info(f"Updating topics display with {len(results_cache)} results")

            self.current_results = dict(results_cache)
            self.update_transcript_list(self.current_results)
            self.root.update()

            logging.info("Topics display update completed")
        except Exception as e:
            logging.error(f"Error updating topics display: {str(e)}", exc_info=True)

    def clear_display(self):
        """Clear all displayed information"""
        self.topic_name_var.set("")
        self.topic_size_var.set("")
        self.topic_coherence_var.set("")

        self.keywords_text.configure(state="normal")
        self.keywords_text.delete(1.0, tk.END)
        self.keywords_text.configure(state="disabled")

        self.chunks_text.configure(state="normal")
        self.chunks_text.delete(1.0, tk.END)
        self.chunks_text.configure(state="disabled")

        self.topics_listbox.delete(0, tk.END)
        self.transcript_combo.set("")

    def show_loading(self, show: bool = True):
        """Show or hide loading state"""
        self.loading = show
        if show:
            self.clear_display()
            self.topic_name_var.set("Loading...")
            self.transcript_combo.configure(state="disabled")
            self.topics_listbox.configure(state="disabled")
        else:
            self.transcript_combo.configure(state="readonly")
            self.topics_listbox.configure(state="normal")
            if self.topic_name_var.get() == "Loading...":
                self.topic_name_var.set("")


class QAExtractorGUI:
    def __init__(self):
        self.tk = tk
        self.filedialog = filedialog
        self.messagebox = messagebox
        self.processor = TranscriptProcessor()
        self.queue = queue.Queue()
        self.custom_words = []
        self.replacements_file_var = None
        self.topics_tab = None  # Add this line
        self.setup_gui()
        self.root.after(100, self.process_queue)

    def setup_gui(self):
        """Setup the GUI components with tabs and enhanced tooltips"""
        self.root = self.tk.Tk()
        self.root.title("Interview Q&A Extractor")
        self.root.geometry("1200x900")

        # Create main container
        main_container = ttk.Frame(self.root, padding="15")
        main_container.grid(row=0, column=0, sticky="nsew")

        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_container.columnconfigure(0, weight=1)
        main_container.rowconfigure(1, weight=1)

        # Create notebook for tabs
        self.notebook = ttk.Notebook(main_container)
        self.topics_tab = TopicsTab(self.notebook, self.queue)
        self.notebook.grid(row=1, column=0, sticky="nsew", pady=10)

        # Create tabs
        self.setup_configuration_tab()
        self.setup_files_tab()
        self.setup_anonymization_tab()
        self.setup_preview_tab()
        self.setup_progress_tab()

        # Set up preview update triggers
        self.transcript_listbox.bind(
            "<<ListboxSelect>>", lambda e: self.refresh_preview_list()
        )

    def create_info_button(self, parent, tooltip_text):
        """Create a circular info button with tooltip"""
        info_button = ttk.Label(
            parent, text="ⓘ", font=("Helvetica", 10), foreground="blue", cursor="hand2"
        )
        self.create_tooltip(info_button, tooltip_text)
        return info_button

    def create_tooltip(self, widget, text):
        """Create a tooltip for a given widget"""

        def show_tooltip(event):
            tooltip = tk.Toplevel()
            tooltip.wm_overrideredirect(True)
            tooltip.wm_geometry(f"+{event.x_root+10}+{event.y_root+10}")

            label = ttk.Label(tooltip, text=text, padding="5")
            label.pack()

            def hide_tooltip(event):
                tooltip.destroy()

            widget.bind("<Leave>", hide_tooltip)
            tooltip.bind("<Leave>", hide_tooltip)

        widget.bind("<Enter>", show_tooltip)

    def setup_configuration_tab(self):
        """Setup the Configuration tab with API settings and processing options"""
        config_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(config_frame, text="Configuration")
        config_frame.columnconfigure(1, weight=1)

        # API Key section
        ttk.Label(
            config_frame, text="OpenAI API Key:", font=("Helvetica", 10, "bold")
        ).grid(row=0, column=0, sticky="w")
        self.api_key_var = self.tk.StringVar()
        key_frame = ttk.Frame(config_frame)
        key_frame.grid(row=0, column=1, sticky="ew", padx=5)
        key_frame.columnconfigure(0, weight=1)

        key_entry = ttk.Entry(key_frame, textvariable=self.api_key_var, show="*")
        key_entry.grid(row=0, column=0, sticky="ew")

        # Key management buttons
        btn_frame = ttk.Frame(key_frame)
        btn_frame.grid(row=0, column=1, padx=(5, 0))
        ttk.Button(
            btn_frame,
            text="📋",
            width=3,
            command=lambda: self.root.clipboard_append(self.api_key_var.get()),
        ).pack(side="left", padx=2)
        ttk.Button(
            btn_frame,
            text="👁",
            width=3,
            command=lambda: key_entry.configure(
                show="" if key_entry.cget("show") else "*"
            ),
        ).pack(side="left", padx=2)

        api_info = self.create_info_button(
            config_frame,
            "Your OpenAI API key is required for processing. Keep it secure and never share it.",
        )
        api_info.grid(row=0, column=2, padx=5)

        # Model selection
        ttk.Label(config_frame, text="GPT Model:", font=("Helvetica", 10, "bold")).grid(
            row=1, column=0, sticky="w", pady=10
        )
        self.model_var = self.tk.StringVar(value="gpt-4o")
        model_combo = ttk.Combobox(
            config_frame,
            textvariable=self.model_var,
            values=["gpt-4o"],
            state="readonly",
        )
        model_combo.grid(row=1, column=1, sticky="w", padx=5)

        model_info = self.create_info_button(
            config_frame,
            "GPT-4 Optimized provides enhanced accuracy for transcript analysis.",
        )
        model_info.grid(row=1, column=2, padx=5)

        # Processing Options
        options_frame = ttk.LabelFrame(
            config_frame, text="Processing Options", padding=10
        )
        options_frame.grid(row=2, column=0, columnspan=3, sticky="ew", pady=10)
        options_frame.columnconfigure(0, weight=1)

        # Create a frame for the checkboxes
        checkbox_frame = ttk.Frame(options_frame)
        checkbox_frame.grid(row=0, column=0, sticky="w", pady=5)

        # Anonymization option
        anon_frame = ttk.Frame(checkbox_frame)
        anon_frame.pack(fill="x", pady=2)
        self.save_anonymized_var = self.tk.BooleanVar(value=False)
        ttk.Checkbutton(
            anon_frame,
            text="Save anonymized transcripts",
            variable=self.save_anonymized_var,
        ).pack(side="left")
        anon_info = self.create_info_button(
            anon_frame,
            "Creates a separate copy of each transcript with sensitive information replaced.",
        )
        anon_info.pack(side="left", padx=5)

        # Sentiment analysis option
        sentiment_frame = ttk.Frame(checkbox_frame)
        sentiment_frame.pack(fill="x", pady=2)
        self.analyze_sentiment_var = self.tk.BooleanVar(value=False)
        ttk.Checkbutton(
            sentiment_frame,
            text="Perform sentiment analysis",
            variable=self.analyze_sentiment_var,
        ).pack(side="left")
        sentiment_info = self.create_info_button(
            sentiment_frame,
            "Analyzes emotional tone in transcript using RoBERTa model.\n"
            "Results are saved in a separate worksheet with '_sent' suffix.\n"
            "Analysis is performed on the original text before anonymization.",
        )
        sentiment_info.pack(side="left", padx=5)

        # Topic analysis option
        topic_frame = ttk.Frame(checkbox_frame)
        topic_frame.pack(fill="x", pady=2)
        self.analyze_topics_var = self.tk.BooleanVar(value=False)
        ttk.Checkbutton(
            topic_frame,
            text="Perform topic analysis",
            variable=self.analyze_topics_var,
        ).pack(side="left")
        topic_info = self.create_info_button(
            topic_frame,
            "Analyzes document topics using BERTopic.\n"
            "Results will be shown in the Topics tab.\n"
            "Analysis is performed locally on the original text.",
        )
        topic_info.pack(side="left", padx=5)

        # Add description text
        desc_text = (
            "Note: Sentiment and topic analysis will be performed securely on your local machine using AI models.\n"
            "GPT analysis will be the only component sent to the cloud, using the anonymized version to protect privacy."
        )
        desc_label = ttk.Label(
            options_frame, text=desc_text, font=("Helvetica", 9, "italic")
        )
        desc_label.grid(row=1, column=0, sticky="w", pady=(5, 0))

    def setup_files_tab(self):
        """Setup the Files tab for file selection"""
        files_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(files_frame, text="Files")
        files_frame.columnconfigure(1, weight=1)

        # Questions File
        ttk.Label(
            files_frame, text="Questions File:", font=("Helvetica", 10, "bold")
        ).grid(row=0, column=0, sticky="w")
        self.questions_file_var = self.tk.StringVar()
        ttk.Entry(files_frame, textvariable=self.questions_file_var).grid(
            row=0, column=1, sticky="ew", padx=5
        )
        ttk.Button(files_frame, text="Browse", command=self.browse_questions).grid(
            row=0, column=2
        )

        questions_info = self.create_info_button(
            files_frame,
            "Select an Excel file containing your questions. The first column should contain one question per row.",
        )
        questions_info.grid(row=0, column=3, padx=5)

        # Transcripts
        ttk.Label(
            files_frame, text="Transcripts:", font=("Helvetica", 10, "bold")
        ).grid(row=1, column=0, sticky="w", pady=10)

        transcript_frame = ttk.Frame(files_frame)
        transcript_frame.grid(row=1, column=1, columnspan=2, sticky="ew", padx=5)
        transcript_frame.columnconfigure(0, weight=1)

        self.transcript_listbox = tk.Listbox(
            transcript_frame, height=6, selectmode=tk.MULTIPLE
        )
        self.transcript_listbox.grid(row=0, column=0, sticky="ew")

        scrollbar = ttk.Scrollbar(
            transcript_frame, orient="vertical", command=self.transcript_listbox.yview
        )
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.transcript_listbox.configure(yscrollcommand=scrollbar.set)

        transcript_info = self.create_info_button(
            files_frame,
            "Add one or more Word document (.docx) transcripts to process.",
        )
        transcript_info.grid(row=1, column=3, padx=5)

        # Transcript buttons
        btn_frame = ttk.Frame(files_frame)
        btn_frame.grid(row=2, column=1, sticky="w", padx=5)
        ttk.Button(btn_frame, text="Add Files", command=self.browse_transcripts).pack(
            side="left", padx=5
        )
        ttk.Button(
            btn_frame, text="Remove Selected", command=self.remove_transcript
        ).pack(side="left")

        # Output File
        ttk.Label(
            files_frame, text="Output File:", font=("Helvetica", 10, "bold")
        ).grid(row=3, column=0, sticky="w", pady=10)
        self.output_file_var = self.tk.StringVar()
        ttk.Entry(files_frame, textvariable=self.output_file_var).grid(
            row=3, column=1, sticky="ew", padx=5
        )
        ttk.Button(files_frame, text="Save As", command=self.browse_output).grid(
            row=3, column=2
        )

        output_info = self.create_info_button(
            files_frame,
            "Select where to save the Excel file containing the analysis results.",
        )
        output_info.grid(row=3, column=3, padx=5)

    def setup_anonymization_tab(self):
        """Setup the Anonymization tab"""
        anon_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(anon_frame, text="Anonymization")
        anon_frame.columnconfigure(0, weight=1)

        # Standard anonymization info
        standard_frame = ttk.LabelFrame(
            anon_frame, text="Standard Anonymization", padding=10
        )
        standard_frame.grid(row=0, column=0, sticky="ew", pady=(0, 10))

        ttk.Label(
            standard_frame,
            text="Standard anonymization replaces the following:",
            font=("Helvetica", 10),
        ).grid(row=0, column=0, sticky="w")

        entities = [
            ("Names & People", "Replaces personal names with [PERSON]"),
            ("Organizations", "Replaces company names with [ORG]"),
            ("Locations", "Replaces cities and locations with [LOC]"),
            ("Dates & Times", "Replaces dates and times with [DATE]/[TIME]"),
            ("Money Values", "Replaces monetary amounts with [MONEY]"),
        ]

        for idx, (entity, tooltip) in enumerate(entities, 1):
            entity_frame = ttk.Frame(standard_frame)
            entity_frame.grid(row=idx, column=0, sticky="w", pady=2)
            ttk.Label(entity_frame, text=f"• {entity}").pack(side="left")
            self.create_info_button(entity_frame, tooltip).pack(side="left", padx=5)

        # Save anonymized copy option
        save_frame = ttk.Frame(standard_frame)
        save_frame.grid(row=len(entities) + 1, column=0, sticky="w", pady=(10, 0))
        self.save_anonymized_var = self.tk.BooleanVar(value=False)
        ttk.Checkbutton(
            save_frame,
            text="Save copy of anonymized transcript",
            variable=self.save_anonymized_var,
        ).pack(side="left")
        save_info = self.create_info_button(
            save_frame,
            "Creates a separate copy of each transcript with anonymization applied.",
        )
        save_info.pack(side="left", padx=5)

        # Custom replacements
        custom_frame = ttk.LabelFrame(
            anon_frame, text="Custom Replacements", padding=10
        )
        custom_frame.grid(row=1, column=0, sticky="ew")

        # Custom word table
        table_frame = ttk.Frame(custom_frame)
        table_frame.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        table_frame.columnconfigure(1, weight=1)

        ttk.Label(table_frame, text="Word to Replace").grid(row=0, column=0, padx=5)
        ttk.Label(table_frame, text="Replacement Text").grid(row=0, column=1, padx=5)

        self.custom_words_frame = ttk.Frame(custom_frame)
        self.custom_words_frame.grid(row=1, column=0, sticky="ew")
        self.custom_words_frame.columnconfigure(1, weight=1)

        # Add/Remove buttons
        btn_frame = ttk.Frame(custom_frame)
        btn_frame.grid(row=2, column=0, pady=10)
        ttk.Button(btn_frame, text="Add Row", command=self.add_replacement_row).pack(
            side="left", padx=5
        )
        ttk.Button(
            btn_frame, text="Remove Row", command=self.remove_replacement_row
        ).pack(side="left")

        # CSV import
        csv_frame = ttk.Frame(custom_frame)
        csv_frame.grid(row=3, column=0, sticky="ew")
        csv_frame.columnconfigure(1, weight=1)

        ttk.Label(csv_frame, text="Import from CSV:").grid(row=0, column=0, sticky="w")
        self.replacements_file_var = self.tk.StringVar()
        ttk.Entry(csv_frame, textvariable=self.replacements_file_var).grid(
            row=0, column=1, padx=5, sticky="ew"
        )
        ttk.Button(csv_frame, text="Browse", command=self.browse_replacements).grid(
            row=0, column=2
        )

        csv_info = self.create_info_button(
            csv_frame,
            "Import replacements from a CSV file. First column: words to replace, Second column: replacement text.",
        )
        csv_info.grid(row=0, column=3, padx=5)

    def setup_preview_tab(self):
        """Setup the Preview tab for anonymization results"""
        preview_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(preview_frame, text="Preview")
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.rowconfigure(1, weight=1)

        # Controls section
        controls_frame = ttk.Frame(preview_frame)
        controls_frame.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        controls_frame.columnconfigure(1, weight=1)

        ttk.Label(controls_frame, text="Select Transcript:").grid(
            row=0, column=0, sticky="w"
        )
        self.preview_transcript_var = self.tk.StringVar()
        self.preview_combo = ttk.Combobox(
            controls_frame, textvariable=self.preview_transcript_var, state="readonly"
        )
        self.preview_combo.grid(row=0, column=1, sticky="ew", padx=5)
        self.preview_combo.bind("<<ComboboxSelected>>", self.update_preview)

        ttk.Button(
            controls_frame, text="↻", width=3, command=self.refresh_preview_list
        ).grid(row=0, column=2)

        preview_info = self.create_info_button(
            controls_frame,
            "Select a transcript to preview how it will look after anonymization.",
        )
        preview_info.grid(row=0, column=3, padx=5)

        # Preview content
        preview_content = ttk.Frame(preview_frame)
        preview_content.grid(row=1, column=0, sticky="nsew")
        preview_content.columnconfigure(0, weight=1)
        preview_content.columnconfigure(1, weight=1)
        preview_content.rowconfigure(1, weight=1)

        # Original text
        ttk.Label(
            preview_content, text="Original Text:", font=("Helvetica", 10, "bold")
        ).grid(row=0, column=0, sticky="w", pady=(0, 5))

        self.original_text = scrolledtext.ScrolledText(
            preview_content, wrap=tk.WORD, width=50, height=30
        )
        self.original_text.grid(row=1, column=0, sticky="nsew", padx=(0, 5))
        self.original_text.configure(state="disabled")

        # Anonymized text
        ttk.Label(
            preview_content, text="Anonymized Text:", font=("Helvetica", 10, "bold")
        ).grid(row=0, column=1, sticky="w", pady=(0, 5))

        self.anonymized_text = scrolledtext.ScrolledText(
            preview_content, wrap=tk.WORD, width=50, height=30
        )
        self.anonymized_text.grid(row=1, column=1, sticky="nsew", padx=(5, 0))
        self.anonymized_text.configure(state="disabled")

        # Statistics section
        stats_frame = ttk.LabelFrame(
            preview_frame, text="Anonymization Statistics", padding=10
        )
        stats_frame.grid(row=2, column=0, sticky="ew", pady=(10, 0))
        stats_frame.columnconfigure(1, weight=1)

        # Statistics labels
        self.stats_labels = {
            "entities": self.tk.StringVar(value="Entities Found: 0"),
            "custom": self.tk.StringVar(value="Custom Words Replaced: 0"),
            "total": self.tk.StringVar(value="Total Replacements: 0"),
        }

        for i, (key, var) in enumerate(self.stats_labels.items()):
            ttk.Label(stats_frame, textvariable=var).grid(
                row=i, column=0, sticky="w", padx=5
            )

    def add_replacement_row(self):
        """Add a new row of entry fields for custom word replacement"""
        row_idx = len(self.custom_words)
        word_var = self.tk.StringVar()
        replacement_var = self.tk.StringVar()

        word_entry = ttk.Entry(self.custom_words_frame, textvariable=word_var)
        replacement_entry = ttk.Entry(
            self.custom_words_frame, textvariable=replacement_var
        )

        word_entry.grid(column=0, row=row_idx + 1, padx=5, pady=2)
        replacement_entry.grid(column=1, row=row_idx + 1, padx=5, pady=2)

        self.custom_words.append((word_var, replacement_var))

    def remove_replacement_row(self):
        """Remove the last row of custom word replacement entries"""
        if self.custom_words:
            last_row = len(self.custom_words)
            for widget in self.custom_words_frame.grid_slaves(row=last_row):
                widget.destroy()
            self.custom_words.pop()

    def browse_replacements(self):
        """Open file dialog for replacements CSV file"""
        filename = self.filedialog.askopenfilename(
            title="Select Replacements CSV File", filetypes=[("CSV files", "*.csv")]
        )
        if filename:
            self.replacements_file_var.set(filename)

    def get_custom_replacements(self) -> Dict[str, str]:
        """Get all custom word replacements from GUI and CSV file"""
        replacements = {}

        # Get replacements from GUI entries
        for word_var, replacement_var in self.custom_words:
            word = word_var.get().strip().lower()
            replacement = replacement_var.get().strip()
            if word and replacement:
                replacements[word] = replacement

        # Get replacements from CSV if specified
        csv_file = self.replacements_file_var.get()
        if csv_file:
            try:
                csv_replacements = self.processor.load_custom_replacements(csv_file)
                replacements.update(csv_replacements)
            except Exception as e:
                self.messagebox.showwarning(
                    "CSV Import Warning",
                    f"Failed to load replacements from CSV: {str(e)}",
                )

        return replacements

    def browse_questions(self):
        """Open file dialog for questions file"""
        filename = self.filedialog.askopenfilename(
            title="Select Questions Excel File", filetypes=[("Excel files", "*.xlsx")]
        )
        if filename:
            self.questions_file_var.set(filename)

    def browse_transcripts(self):
        """Open file dialog for transcript files"""
        filenames = self.filedialog.askopenfilenames(
            title="Select Transcript Files", filetypes=[("Word files", "*.docx")]
        )
        if filenames:
            for filename in filenames:
                if filename not in self.transcript_listbox.get(0, self.tk.END):
                    self.transcript_listbox.insert(self.tk.END, filename)

    def browse_output(self):
        """Open file dialog for output file"""
        filename = self.filedialog.asksaveasfilename(
            title="Save Output Excel File",
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx")],
        )
        if filename:
            self.output_file_var.set(filename)

    def remove_transcript(self):
        """Remove selected transcript from the list"""
        selections = self.transcript_listbox.curselection()
        if selections:
            for index in reversed(selections):
                self.transcript_listbox.delete(index)

    def refresh_preview_list(self):
        """Update the preview dropdown with current transcript list"""
        transcripts = list(self.transcript_listbox.get(0, self.tk.END))
        self.preview_combo["values"] = transcripts
        if transcripts and not self.preview_transcript_var.get():
            self.preview_combo.set(transcripts[0])
            self.update_preview(None)

    def update_preview(self, event=None):
        """Update the preview texts when a transcript is selected"""
        selected_file = self.preview_transcript_var.get()
        if not selected_file:
            return

        try:
            # Load and display original text
            original = self.processor.load_transcript(selected_file)
            self.original_text.configure(state="normal")
            self.original_text.delete(1.0, tk.END)
            self.original_text.insert(1.0, original)
            self.original_text.configure(state="disabled")

            # Get custom replacements
            custom_replacements = self.get_custom_replacements()

            # Generate and display anonymized text
            anonymized = self.processor.anonymize_text(original, custom_replacements)
            self.anonymized_text.configure(state="normal")
            self.anonymized_text.delete(1.0, tk.END)
            self.anonymized_text.insert(1.0, anonymized)
            self.anonymized_text.configure(state="disabled")

            # Update statistics
            self._update_statistics(original, anonymized, custom_replacements)

        except Exception as e:
            self.messagebox.showerror(
                "Preview Error", f"Error generating preview:\n{str(e)}"
            )

    def _update_statistics(self, original, anonymized, custom_replacements):
        """Update the anonymization statistics"""
        # Count standard entities
        entity_count = len(
            re.findall(r"\[(PERSON|ORG|GPE|LOC|DATE|TIME|MONEY)\]", anonymized)
        )

        # Count custom word replacements
        custom_count = 0
        if custom_replacements:
            for original_word in custom_replacements.keys():
                custom_count += len(
                    re.findall(
                        r"\b" + re.escape(original_word) + r"\b",
                        original,
                        re.IGNORECASE,
                    )
                )

        # Update statistics labels
        self.stats_labels["entities"].set(f"Entities Found: {entity_count}")
        self.stats_labels["custom"].set(f"Custom Words Replaced: {custom_count}")
        self.stats_labels["total"].set(
            f"Total Replacements: {entity_count + custom_count}"
        )

    def setup_progress_tab(self):
        """Setup the Progress tab with individual transcript progress tracking and time estimates"""
        progress_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(progress_frame, text="Progress")
        progress_frame.columnconfigure(0, weight=1)
        progress_frame.rowconfigure(1, weight=1)

        # Status section
        status_frame = ttk.LabelFrame(progress_frame, text="Overall Status", padding=10)
        status_frame.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        status_frame.columnconfigure(0, weight=1)

        # Status and time estimate
        status_info_frame = ttk.Frame(status_frame)
        status_info_frame.grid(row=0, column=0, sticky="ew")
        status_info_frame.columnconfigure(1, weight=1)

        self.status_var = self.tk.StringVar(value="Ready")
        ttk.Label(
            status_info_frame,
            textvariable=self.status_var,
            font=("Helvetica", 10, "bold"),
        ).grid(row=0, column=0, sticky="w")

        self.overall_time_var = self.tk.StringVar()
        ttk.Label(
            status_info_frame,
            textvariable=self.overall_time_var,
            font=("Helvetica", 10),
        ).grid(row=0, column=1, sticky="e")

        # Overall progress
        overall_progress_frame = ttk.Frame(status_frame)
        overall_progress_frame.grid(row=1, column=0, sticky="ew", pady=(10, 0))
        overall_progress_frame.columnconfigure(0, weight=1)

        self.progress_var = self.tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            overall_progress_frame, variable=self.progress_var, mode="determinate"
        )
        self.progress_bar.grid(row=0, column=0, sticky="ew", padx=(0, 5))
        self.progress_label = ttk.Label(overall_progress_frame, text="0%")
        self.progress_label.grid(row=0, column=1)

        # Transcript progress section
        transcript_frame = ttk.LabelFrame(
            progress_frame, text="Transcript Progress", padding=10
        )
        transcript_frame.grid(row=1, column=0, sticky="nsew")
        transcript_frame.columnconfigure(0, weight=1)
        transcript_frame.rowconfigure(0, weight=1)

        # Create scrollable frame for transcript progress
        canvas = tk.Canvas(transcript_frame)
        scrollbar = ttk.Scrollbar(
            transcript_frame, orient="vertical", command=canvas.yview
        )
        self.transcript_progress_frame = ttk.Frame(canvas)

        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")
        canvas.create_window((0, 0), window=self.transcript_progress_frame, anchor="nw")

        self.transcript_progress_frame.bind(
            "<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        self.transcript_progress_frame.columnconfigure(1, weight=1)

        # Initialize progress tracking dictionaries
        self.transcript_progress = {}
        self.time_estimators = {}
        self.overall_estimator = TimeEstimator()

        # Current operation
        self.question_progress_var = self.tk.StringVar()
        ttk.Label(
            status_frame, textvariable=self.question_progress_var, wraplength=800
        ).grid(row=2, column=0, sticky="w", pady=(10, 0))

        # Log viewer
        log_frame = ttk.LabelFrame(progress_frame, text="Log", padding=10)
        log_frame.grid(row=2, column=0, sticky="ew", pady=(10, 0))
        log_frame.columnconfigure(0, weight=1)

        self.log_text = scrolledtext.ScrolledText(
            log_frame, height=10, width=90, wrap=tk.WORD
        )
        self.log_text.grid(row=0, column=0, sticky="ew")
        self.log_text.configure(state="disabled")

        # Start button
        self.process_button = ttk.Button(
            progress_frame,
            text="Start Processing",
            command=self.process_transcripts,
            style="Accent.TButton",
        )
        self.process_button.grid(row=3, column=0, pady=10)

    def initialize_transcript_progress(self):
        """Initialize progress tracking for each transcript with time estimates"""
        # Clear existing progress widgets
        for widget in self.transcript_progress_frame.winfo_children():
            widget.destroy()

        # Reset tracking dictionaries
        self.transcript_progress = {}
        self.time_estimators = {}

        # Initialize overall time estimator
        transcripts = list(self.transcript_listbox.get(0, tk.END))
        total_questions = len(
            self.processor.load_questions(self.questions_file_var.get())
        )
        self.overall_estimator.start(len(transcripts) * total_questions)

        # Create progress bars for each transcript
        for idx, transcript in enumerate(transcripts):
            frame = ttk.Frame(self.transcript_progress_frame)
            frame.grid(row=idx, column=0, sticky="ew", pady=2, padx=5)
            frame.columnconfigure(1, weight=1)

            # Transcript name
            name_label = ttk.Label(
                frame, text=os.path.basename(transcript), width=30, anchor="w"
            )
            name_label.grid(row=0, column=0, padx=(0, 10))

            # Progress bar
            progress_var = self.tk.DoubleVar()
            progress_bar = ttk.Progressbar(
                frame, variable=progress_var, mode="determinate"
            )
            progress_bar.grid(row=0, column=1, sticky="ew")

            # Status and time estimate
            info_frame = ttk.Frame(frame)
            info_frame.grid(row=0, column=2, padx=(10, 0))

            status_var = self.tk.StringVar(value="Pending")
            status_label = ttk.Label(
                info_frame, textvariable=status_var, width=10, anchor="center"
            )
            status_label.grid(row=0, column=0, padx=(0, 5))

            time_var = self.tk.StringVar()
            time_label = ttk.Label(
                info_frame, textvariable=time_var, width=20, anchor="e"
            )
            time_label.grid(row=0, column=1)

            # Store progress tracking info
            self.transcript_progress[transcript] = {
                "progress_var": progress_var,
                "status_var": status_var,
                "time_var": time_var,
            }

            # Initialize time estimator for this transcript
            self.time_estimators[transcript] = TimeEstimator()
            self.time_estimators[transcript].start(total_questions)

    def update_transcript_progress(self, transcript, progress=None, status=None):
        """Update progress, status, and time estimates for a specific transcript"""
        if transcript in self.transcript_progress:
            if progress is not None:
                self.transcript_progress[transcript]["progress_var"].set(progress)

                # Update time estimates
                estimator = self.time_estimators[transcript]
                estimator.update(progress)
                time_estimate = estimator.get_estimate(progress)
                if time_estimate:
                    self.transcript_progress[transcript]["time_var"].set(
                        f"Remaining: {time_estimate}"
                    )
                elif progress >= 100:
                    self.transcript_progress[transcript]["time_var"].set("Completed")

            if status is not None:
                self.transcript_progress[transcript]["status_var"].set(status)
                if status == "Error":
                    self.transcript_progress[transcript]["time_var"].set("Failed")

    def process_queue(self):
        """Process queue messages including time estimate updates"""
        try:
            while True:
                message = self.queue.get_nowait()
                msg_type, content = message

                if msg_type == "topics_update":
                    if self.topics_tab:
                        logging.info("Received topics update with content:")
                        logging.info(f"Number of transcripts: {len(content)}")
                        logging.info(f"Content keys: {list(content.keys())}")
                        self.topics_tab.update_results(content)
                        self.root.update()
                    else:
                        logging.warning("Topics tab not initialized")

                elif msg_type == "topics_loading":
                    if self.topics_tab:
                        self.topics_tab.show_loading(content)

                elif msg_type == "transcript_progress":
                    transcript, progress = content
                    self.update_transcript_progress(transcript, progress=progress)
                    self.overall_estimator.update(self.progress_var.get())
                    overall_estimate = self.overall_estimator.get_estimate(
                        self.progress_var.get()
                    )
                    if overall_estimate:
                        self.overall_time_var.set(
                            f"Estimated time remaining: {overall_estimate}"
                        )

                elif msg_type == "transcript_status":
                    transcript, status = content
                    self.update_transcript_progress(transcript, status=status)

                elif msg_type == "log":
                    self.log_text.configure(state="normal")
                    self.log_text.insert(self.tk.END, content)
                    self.log_text.see(self.tk.END)
                    self.log_text.configure(state="disabled")

                elif msg_type == "progress":
                    self.progress_var.set(content)
                    self.progress_label.configure(text=f"{int(content)}%")

                elif msg_type == "status":
                    self.status_var.set(content)

                elif msg_type == "question_progress":
                    self.question_progress_var.set(content)

                elif msg_type == "info":
                    self.messagebox.showinfo("Success", content)

                elif msg_type == "error_prompt":
                    continue_processing = self.messagebox.askyesno("Error", content)
                    if not continue_processing:
                        self.queue.put(("done", None))
                        return

                elif msg_type == "error":
                    self.messagebox.showerror("Error", content)

                elif msg_type == "done":
                    self.process_button.configure(state="normal")
                    if self.status_var.get() != "Completed":
                        self.status_var.set("Ready")
                    self.question_progress_var.set("")
                    self.overall_time_var.set("")

        except queue.Empty:
            pass
        finally:
            self.root.after(100, self.process_queue)

    def validate_inputs(self):
        """Validate all required inputs before processing"""
        if not self.api_key_var.get():
            self.messagebox.showerror("Error", "Please enter your OpenAI API key")
            return False

        if not self.questions_file_var.get():
            self.messagebox.showerror("Error", "Please select a questions file")
            return False

        if self.transcript_listbox.size() == 0:
            self.messagebox.showerror(
                "Error", "Please select at least one transcript file"
            )
            return False

        if not self.output_file_var.get():
            self.messagebox.showerror("Error", "Please select an output file location")
            return False

        return True

    def process_transcripts(self):
        """Start processing transcripts with sentiment analysis option"""
        if not self.validate_inputs():
            return

        try:
            # Disable the process button
            self.process_button.configure(state="disabled")
            self.question_progress_var.set("")

            # Initialize progress tracking
            self.initialize_transcript_progress()

            # Update status
            self.queue.put(("status", "Initializing processing..."))
            self.queue.put(("log", "Starting transcript processing...\n"))

            if self.analyze_sentiment_var.get():
                self.queue.put(("log", "Sentiment analysis is enabled.\n"))

            # Start processing thread
            threading.Thread(
                target=self._process_transcripts_thread,
                kwargs={"analyze_sentiment": self.analyze_sentiment_var.get()},
                daemon=True,
            ).start()

        except Exception as e:
            error_msg = f"Failed to start processing: {str(e)}"
            self.queue.put(("log", f"\nERROR: {error_msg}\n"))
            self.queue.put(("error", error_msg))
            self.process_button.configure(state="normal")
            logging.error("Error starting processing", exc_info=True)

    def _process_transcripts_thread(self, analyze_sentiment: bool = False):
        """Process transcripts with sentiment analysis"""
        try:
            # Initialize processor and load questions
            custom_replacements = self.get_custom_replacements()
            self.queue.put(("status", "Setting API key..."))
            self.processor.set_api_key(self.api_key_var.get())
            self.processor.model_name = self.model_var.get()

            self.queue.put(("log", "Loading questions...\n"))
            questions = self.processor.load_questions(self.questions_file_var.get())
            self.queue.put(("log", f"Loaded {len(questions)} questions\n"))

            transcripts = list(self.transcript_listbox.get(0, tk.END))
            total_transcripts = len(transcripts)

            if self.analyze_topics_var.get():
                self.queue.put(("topics_loading", True))

            # Calculate total steps (questions per transcript + sentiment if enabled)
            steps_per_transcript = len(questions)
            if analyze_sentiment:
                steps_per_transcript += 1
            total_steps = total_transcripts * steps_per_transcript

            completed_steps = 0

            for i, transcript_file in enumerate(transcripts, 1):
                try:
                    # Update overall progress
                    base_progress = ((i - 1) / total_transcripts) * 100
                    self.queue.put(("progress", base_progress))

                    # Update transcript status
                    self.queue.put(
                        ("transcript_status", (transcript_file, "Processing"))
                    )
                    self.queue.put(
                        (
                            "log",
                            f"\nProcessing transcript {i}/{total_transcripts}:"
                            f"\n{transcript_file}\n"
                            "----------------------------------------\n",
                        )
                    )

                    def progress_callback(msg_type, msg_content):
                        nonlocal completed_steps
                        if msg_type == "question_progress":
                            try:
                                # Update progress based on current step
                                if "question" in msg_content.lower():
                                    match = re.search(r"(\d+)/\d+", msg_content)
                                    if match:
                                        question_num = int(match.group(1))
                                        progress = (
                                            (
                                                (i - 1) * steps_per_transcript
                                                + question_num
                                            )
                                            / total_steps
                                        ) * 100

                                        # Update individual transcript progress
                                        transcript_progress = (
                                            question_num / len(questions)
                                        ) * 100
                                        self.queue.put(
                                            (
                                                "transcript_progress",
                                                (transcript_file, transcript_progress),
                                            )
                                        )

                                        # Update overall progress
                                        self.queue.put(("progress", progress))
                                elif "sentiment" in msg_content.lower():
                                    progress = (
                                        (
                                            (i - 1) * steps_per_transcript
                                            + completed_steps
                                        )
                                        / total_steps
                                    ) * 100
                                    self.queue.put(("progress", progress))

                            except (ValueError, AttributeError) as e:
                                logging.warning(f"Error parsing progress: {e}")

                            self.queue.put(("question_progress", msg_content))
                            self.queue.put(("log", f"{msg_content}\n"))

                        elif msg_type == "status":
                            self.queue.put(("status", msg_content))

                    # Process transcript with analysis options
                    answers = self.processor.process_single_transcript(
                        transcript_file,
                        questions,
                        anonymize=self.save_anonymized_var.get(),
                        progress_callback=progress_callback,
                        custom_words=custom_replacements,
                        analyze_sentiment=analyze_sentiment,
                        analyze_topics=self.analyze_topics_var.get(),
                    )

                    # Update topic results if needed
                    if self.analyze_topics_var.get():
                        self.queue.put(
                            (
                                "topics_update",
                                self.processor.topic_analyzer.results_cache,
                            )
                        )

                        # Save results
                        self.queue.put(("log", "Saving results to Excel...\n"))
                        self.processor.save_results(
                            answers,
                            questions,
                            self.output_file_var.get(),
                            transcript_file,
                        )

                        # Update completion status
                        self.queue.put(
                            ("transcript_status", (transcript_file, "Completed"))
                        )
                        self.queue.put(("transcript_progress", (transcript_file, 100)))

                        # Log completion details
                        completion_msg = f"Complete! Results saved to sheet '{os.path.basename(transcript_file)}'"
                        self.queue.put(("log", f"{completion_msg}\n"))

                        completed_steps += steps_per_transcript

                except Exception as e:
                    error_msg = (
                        f"Error processing transcript {transcript_file}:\n{str(e)}\n"
                    )
                    self.queue.put(("log", f"\nERROR: {error_msg}\n"))
                    self.queue.put(("transcript_status", (transcript_file, "Error")))
                    logging.error(error_msg, exc_info=True)

                    if i < total_transcripts:
                        continue_msg = (
                            f"Error processing {os.path.basename(transcript_file)}.\n"
                            "Would you like to continue with the remaining transcripts?"
                        )
                        self.queue.put(("error_prompt", continue_msg))
                        try:
                            response = self.queue.get(
                                timeout=5
                            )  # Added timeout to prevent blocking
                            if response[0] == "done":
                                break
                        except queue.Empty:
                            logging.warning("No response received for error prompt.")
                            break
                    else:
                        self.queue.put(
                            (
                                "error",
                                f"Failed to process {os.path.basename(transcript_file)}",
                            )
                        )

            # Complete
            if self.analyze_topics_var.get():
                self.queue.put(("topics_loading", False))

            self.queue.put(("progress", 100))
            self.queue.put(("question_progress", ""))
            self.queue.put(("log", "\nAll processing completed!\n"))
            self.queue.put(("status", "Completed"))

            # Show completion message with sentiment analysis info
            completion_msg = "Processing completed successfully!"
            if analyze_sentiment:
                completion_msg += (
                    "\nSentiment analysis results are included in the answers."
                )
            self.queue.put(("info", completion_msg))

        except Exception as e:
            error_msg = f"An error occurred: {str(e)}"
            self.queue.put(("log", f"\nError: {error_msg}\n"))
            self.queue.put(("error", error_msg))
            logging.error("Processing error", exc_info=True)

        finally:
            self.queue.put(("done", None))


def main():
    """Main entry point for the application with proper resource cleanup"""
    # Initialize multiprocessing cleanup handler
    import atexit
    import multiprocessing

    # Create a list to track resources
    active_processes = []

    def cleanup_resources():
        """Cleanup function to handle multiprocessing resources"""
        # Clean up any active processes
        for process in active_processes:
            if process.is_alive():
                process.terminate()
                process.join()

        # Shutdown multiprocessing
        if hasattr(multiprocessing, "_resource_tracker"):
            multiprocessing._resource_tracker._resource_tracker.clear()
            multiprocessing._resource_tracker._resource_tracker._stop_event.set()
            if hasattr(
                multiprocessing._resource_tracker._resource_tracker, "_check_process"
            ):
                multiprocessing._resource_tracker._resource_tracker._check_process.join()

        # Clean up any remaining semaphores
        try:
            multiprocessing.resource_tracker._resource_tracker.clear()
        except Exception:
            pass

    # Register the cleanup function
    atexit.register(cleanup_resources)

    try:
        # Check Python version and virtual environment
        if not is_python312():
            reexecute_with_python312()

        if not is_venv_active():
            if not os.path.exists(VENV_DIR):
                create_virtual_env(VENV_DIR)
            reexecute_with_venv(VENV_DIR)

        # Setup all dependencies
        setup_dependencies()

        # Start the GUI application
        app = QAExtractorGUI()

        # Add cleanup method to the GUI class
        def on_closing():
            """Handle window closing event"""
            cleanup_resources()
            app.root.destroy()

        app.root.protocol("WM_DELETE_WINDOW", on_closing)
        app.root.mainloop()

    except Exception as e:
        logging.critical("Unhandled exception in main: %s", e, exc_info=True)
        messagebox.showerror(
            "Critical Error",
            f"A critical error occurred:\n{e}\nCheck the log file for more details.",
        )
        cleanup_resources()
        sys.exit(1)

    finally:
        # Ensure cleanup runs even if an error occurs
        cleanup_resources()


if __name__ == "__main__":
    main()
